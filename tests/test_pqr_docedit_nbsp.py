# -*- coding: utf-8 -*-
"""문구 찾기가 빈칸 종류에 걸려 넘어지지 않아야 한다.

디겐타안연고 결재본은 'QC-126\xa0제품품질평가규정' 처럼 사이가 줄바꿈 없는 공백(\xa0)이라,
퀴노비드 기준의 'QC-126 제품품질평가규정' 으로 찾지 못해 엔진이 9항 각주에서 멈췄다
(2026-09). 한 글자도 다르지 않은데 빈칸 종류만 달랐고, 그래서 보고서에 수치가 하나도
들어가지 않았다.
"""
import unittest

import docx

from pqr.engine import docedit as E


def _doc(*texts):
    d = docx.Document()
    for t in texts:
        d.add_paragraph(t)
    return d


class 빈칸을_가리지_않는다(unittest.TestCase):
    def test_줄바꿈_없는_공백도_찾는다(self):
        d = _doc("- QC-126 제품품질평가규정")
        self.assertIsNotNone(E.find_para(d, "QC-126 제품품질평가규정"))

    def test_일반_공백은_그대로_찾는다(self):
        d = _doc("- QC-126 제품품질평가규정")
        self.assertIsNotNone(E.find_para(d, "QC-126 제품품질평가규정"))

    def test_없는_문구는_없다고_한다(self):
        self.assertIsNone(E.find_para(_doc("다른 글"), "QC-126 제품품질평가규정"))

    def test_빈칸_수가_달라도_찾는다(self):
        self.assertIsNotNone(E.find_para(_doc("16.  결   론"), "16. 결 론"))

    def test_바꾸기도_빈칸을_가리지_않는다(self):
        d = _doc("제품품질평가는 2024년 1월 ~ 12월까지 생산된 제품을 평가한다.")
        E.replace_in_para(d, "제품품질평가는", "2024년 1월", "2025년 1월")
        self.assertIn("2025년 1월", d.paragraphs[0].text)


class 각주는_본보기가_없어도_넣는다(unittest.TestCase):
    """각주 하나 때문에 보고서 전체를 못 만들면 안 된다."""

    def test_본보기_문단이_없어도_넣는다(self):
        d = docx.Document()
        d.add_paragraph("아무 글")
        table = d.add_table(rows=1, cols=1)
        E.note_after(d, table, "1) 각주입니다.")
        self.assertIn("1) 각주입니다.", [p.text for p in d.paragraphs])

    def test_본보기가_있으면_그_서식을_쓴다(self):
        d = docx.Document()
        sample = d.add_paragraph("- QC-126 제품품질평가규정")
        sample.runs[0].bold = True
        table = d.add_table(rows=1, cols=1)
        E.note_after(d, table, "2) 각주입니다.")
        made = next(p for p in d.paragraphs if p.text == "2) 각주입니다.")
        self.assertTrue(made.runs[0].bold)


if __name__ == "__main__":
    unittest.main()


class 주성분이_둘_이상이면(unittest.TestCase):
    """디겐타안연고는 플루오로메톨론·겐타마이신황산염 두 성분이라 9.1 에 함량 줄이 둘이다.

    성적서에서 읽은 함량은 한 벌뿐이라 같은 값이 두 줄에 들어간다. 어느 값이 어느 성분
    것인지는 원본을 봐야 안다 — 지어내지 않고 짚어 준다.
    """

    def test_허용기준_앞머리의_성분_이름을_읽는다(self):
        from pqr.engine.recipe_ointment import assay_component
        self.assertEqual(assay_component("[플루오로메톨론]\n허가) 90.0 ~ 110.0 %"), "플루오로메톨론")
        self.assertEqual(assay_component(" [겐타마이신황산염] 허가) 90.0 ~ 120.0 %"), "겐타마이신황산염")

    def test_성분_표시가_없으면_없다고_한다(self):
        from pqr.engine.recipe_ointment import assay_component
        self.assertIsNone(assay_component("허가) 90.0 ~ 110.0 %"))
        self.assertIsNone(assay_component(""))
        self.assertIsNone(assay_component("허가) 90.0 % [주1]"))    # 앞머리가 아니면 성분이 아니다
