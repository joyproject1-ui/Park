# -*- coding: utf-8 -*-
"""9.2 소제목 번호와 질량 각주 — 담당자 2026-09 점검(디겐타안연고 2026 양식).

양식에 '9.2.1 · 9.2.2 · 9.2.1.3 포장 완료 후' 처럼 번호가 어긋나 있고, 질량 각주가 양식의
'1) 모든 Lot …' 각주와 같은 번호로 다른 표 밑에 붙어 있었다."""
from __future__ import unicode_literals

import unittest

import docx

from pqr.engine import docedit as E
from pqr.engine import recipe_ointment as R


def doc_with(*lines):
    d = docx.Document()
    for line in lines:
        d.add_paragraph(line)
    return d


class 소제목_번호(unittest.TestCase):
    def test_어긋난_셋째_번호를_바로잡는다(self):
        d = doc_with("9. 중요공정", "9.2 세부 시험결과", "9.2.1 조제 완료 후", "9.2.2 충전 완료 후",
                     "9.2.1.3 포장 완료 후", "10. 밸리데이션 현황표")
        self.assertEqual(E.renumber_subheadings(d, "9.2"), 1)
        self.assertEqual([p.text for p in d.paragraphs][4], "9.2.3 포장 완료 후")

    def test_정상적인_하위_번호는_두다(self):
        d = doc_with("9.2 세부 시험결과", "9.2.1 조제 완료 후", "9.2.1.1 성상표", "9.2.1.2 이물표",
                     "9.2.2 충전 완료 후", "9.2.2.1 질량표")
        self.assertEqual(E.renumber_subheadings(d, "9.2"), 0)

    def test_다른_항의_번호는_건드리지_않는다(self):
        d = doc_with("9.2.1 조제", "10.2 제조설비", "10.2.1 어떤 표")
        self.assertEqual(E.renumber_subheadings(d, "9.2"), 0)


class 각주_번호(unittest.TestCase):
    def test_양식에_이미_있는_각주_번호를_센다(self):
        d = doc_with("9.2 세부 시험결과", "9.2.1.3 포장 완료 후",
                     "1) 모든 Lot 의 시험결과가 0 개(매)로 동일하여 최댓값, 최솟값, 평균은 별도로 작성하지 않음.",
                     "10. 밸리데이션 현황표", "1) 다른 항의 각주")
        self.assertEqual(R._note_numbers_under(d, "9.2"), {1})

    def test_머리행_칸에_번호를_붙인다(self):
        d = docx.Document()
        t = d.add_table(rows=2, cols=3)
        t.cell(0, 0).text = "연번"; t.cell(0, 1).text = "평균"; t.cell(0, 2).text = "개개"
        self.assertTrue(R._mark_header(t, "개개", 2))
        self.assertEqual(t.cell(0, 2).text, "개개2)")
        self.assertFalse(R._mark_header(t, "없는말", 2))


if __name__ == "__main__":
    unittest.main()
