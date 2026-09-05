# -*- coding: utf-8 -*-
"""rPr·pPr 이 없는 run·문단에 속성 요소를 만들어 넣을 때 — 담당자 PC 양식(최종 공양식 2026PQR)의
12항 run 에 rPr 이 없어 KeyError 'r' 로 보고서가 요약본으로 떨어졌다(2026-09-05 작성 기록)."""
from __future__ import unicode_literals

import unittest

import docx
from docx.oxml.ns import qn

from pqr.engine import docedit as E
from pqr.engine.ooxml_order import get_or_add, place, resort


class 속성요소_없는_run(unittest.TestCase):
    def bare_run(self):
        d = docx.Document()
        p = d.add_paragraph()
        r = p.add_run("[CC-240423-06] 변경 건\n내용 줄")
        rpr = r._r.find(qn("w:rPr"))
        if rpr is not None:
            r._r.remove(rpr)
        ppr = p._p.find(qn("w:pPr"))
        if ppr is not None:
            p._p.remove(ppr)
        return d, p, r

    def test_rPr_은_run_의_맨_앞에_생긴다(self):
        _, _, r = self.bare_run()
        rpr = get_or_add(r._r, "rPr")
        self.assertIs(r._r[0], rpr)
        self.assertEqual(r._r[1].tag.split("}")[-1], "t")

    def test_pPr_은_문단의_맨_앞에_생긴다(self):
        _, p, _ = self.bare_run()
        ppr = get_or_add(p._p, "pPr")
        self.assertIs(p._p[0], ppr)

    def test_첫_줄만_굵게가_넘어지지_않는다(self):
        d = docx.Document()
        t = d.add_table(rows=1, cols=2)
        cell = t.cell(0, 1)
        r = cell.paragraphs[0].add_run("[CC-240423-06] 변경 건")
        rpr = r._r.find(qn("w:rPr"))
        if rpr is not None:
            r._r.remove(rpr)
        second = cell.add_paragraph()
        r2 = second.add_run("아래 내용"); r2.bold = True
        E.bold_first_line_only(cell)                      # 예전엔 KeyError: 'r'
        self.assertIsNot(r._r[0].tag.split("}")[-1] == "rPr", False)
        self.assertFalse(cell.paragraphs[1].runs[0].bold)   # 아래 줄은 보통으로

    def test_모르는_부모면_뒤에_붙이고_resort_는_그냥_둔다(self):
        d = docx.Document()
        p = d.add_paragraph("x")
        body = p._p.getparent()
        el = body.makeelement(qn("w:bookmarkEnd"), {})
        place(body, el)
        self.assertIs(body[-1], el)
        self.assertFalse(resort(body))


if __name__ == "__main__":
    unittest.main()
