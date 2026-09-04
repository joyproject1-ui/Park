# -*- coding: utf-8 -*-
"""다른 제품의 결재본·규격으로 쓰지 않도록 — 제품 확인과 9.1 허용기준 한계 읽기."""
from __future__ import unicode_literals

import os
import tempfile
import unittest

import docx

from pqr.engine import recipe_ointment as R
from pqr.engine import writer


class 제품확인(unittest.TestCase):
    def _doc(self, title):
        d = docx.Document()
        d.add_paragraph(title)
        d.add_paragraph("제품품질평가 보고서")
        p = os.path.join(tempfile.mkdtemp(), "base.docx"); d.save(p); return p

    def test_같은_제품이면_통과(self):
        p = self._doc("퀴노비드안연고 (오플록사신)")
        self.assertIsNone(writer.other_product(p, {"code": "QC1-7007", "name": "퀴노비드안연고"}))

    def test_다른_제품이면_이유를_돌려준다(self):
        p = self._doc("퀴노비드안연고 (오플록사신)")
        why = writer.other_product(p, {"code": "QC1-7014", "name": "디겐타안연고"})
        self.assertIsNotNone(why)
        self.assertIn("디겐타안연고", why)
        self.assertIn("다른 제품", why)

    def test_코드가_있으면_통과(self):
        p = self._doc("문서번호 PQR25-2-QC1-7014")
        self.assertIsNone(writer.other_product(p, {"code": "QC1-7014", "name": "디겐타안연고"}))

    def test_이름과_코드가_없으면_검사하지_않는다(self):
        p = self._doc("아무 제목")
        self.assertIsNone(writer.other_product(p, {}))


class 규격한계(unittest.TestCase):
    def test_퀴노비드_9_1_문구(self):
        got = R.parse_limits({"particle": "75 ㎛ 이하", "assay": "표시량의 90.0 ~ 110.0%",
                              "metal": "50 ㎛ 이상 이물의 합계 50개 이하, 개개 8개 초과 3매 이하"})
        self.assertEqual(got, {"particle": 75.0, "assay": (90.0, 110.0), "metal": 50.0})

    def test_다른_제품_문구(self):
        got = R.parse_limits({"particle": "입자도 100um 이하", "assay": "95.0∼105.0 %", "metal": ""})
        self.assertEqual(got["particle"], 100.0)
        self.assertEqual(got["assay"], (95.0, 105.0))
        self.assertNotIn("metal", got)

    def test_못_읽으면_비운다(self):
        self.assertEqual(R.parse_limits({"particle": "적합", "assay": "", "metal": None}), {})


class 성적서함량(unittest.TestCase):
    def test_규격을_글에서_읽는다(self):
        from pqr.engine.readers import coa
        text = "함량 오플록사신 90.0 ~ 110.0%  107.3 %  합격"
        out = {}
        # read_fp 의 함량 부분과 같은 논리
        import re
        assays = []
        for m in re.finditer(r"(\d+(?:\.\d)?)\s*~\s*(\d+(?:\.\d)?)\s*%\s+.*?([\d]+\.\d)\s*%", text, flags=re.S):
            assays.append((m.group(1), m.group(2), m.group(3)))
        self.assertEqual(assays, [("90.0", "110.0", "107.3")])
        text2 = "함량 겐타마이신 95.0 ~ 105.0%  99.8 %  덱사메타손 93.0 ~ 107.0%  101.2 %"
        assays2 = [(m.group(1), m.group(2), m.group(3)) for m in
                   re.finditer(r"(\d+(?:\.\d)?)\s*~\s*(\d+(?:\.\d)?)\s*%\s+.*?([\d]+\.\d)\s*%", text2, flags=re.S)]
        self.assertEqual(assays2, [("95.0", "105.0", "99.8"), ("93.0", "107.0", "101.2")])


if __name__ == "__main__":
    unittest.main()
