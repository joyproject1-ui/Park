# -*- coding: utf-8 -*-
"""담당자가 결재본을 보며 짚어 준 차림새 규칙들 (2026-09, 디겐타안연고).

  · 사선을 그은 칸의 'N/A' 는 지운다 — 사선이 곧 해당 없음이다
  · '확인 필요' 는 노랑 형광펜으로 — 직접 채워야 할 칸이 한눈에 보이게
  · 각주 줄과 다음 항 제목 사이의 빈 줄은 지운다
  · 변경관리 변경사항은 큰 항목만, 조치사항에서 변경통보는 뺀다
"""
import unittest

import docx

from pqr.engine import docedit as E
from pqr.engine.recipe_ointment import brief_change


class 사선_칸의_NA(unittest.TestCase):
    def _doc(self, text, diag):
        document = docx.Document()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "비고"
        table.cell(0, 1).text = text
        if diag:
            E.add_diag(table.cell(0, 1))
        return document, table

    def test_사선이_있으면_지운다(self):
        document, table = self._doc("N/A", True)
        self.assertEqual(E.drop_na_in_diag_cells(document), 1)
        self.assertEqual(table.cell(0, 1).text.strip(), "")

    def test_사선이_없으면_둔다(self):
        document, table = self._doc("N/A", False)
        self.assertEqual(E.drop_na_in_diag_cells(document), 0)
        self.assertEqual(table.cell(0, 1).text.strip(), "N/A")

    def test_다른_글은_지우지_않는다(self):
        document, table = self._doc("해당 Lot 없음", True)
        self.assertEqual(E.drop_na_in_diag_cells(document), 0)


class 확인_필요_표시(unittest.TestCase):
    def test_그_런에만_형광펜(self):
        document = docx.Document()
        document.add_paragraph("확인 필요")
        document.add_paragraph("적합")
        self.assertEqual(E.highlight(document, "확인 필요"), 1)
        runs = document.paragraphs[0].runs
        self.assertEqual(runs[0].font.highlight_color, docx.enum.text.WD_COLOR_INDEX.YELLOW)
        self.assertIsNone(document.paragraphs[1].runs[0].font.highlight_color)


class 각주_뒤_빈_줄(unittest.TestCase):
    def _doc(self, lines):
        document = docx.Document()
        for text in lines:
            document.add_paragraph(text)
        return document

    def test_각주와_항_제목_사이는_지운다(self):
        document = self._doc(["* [RSF102] 플루오로메톨론 사용이력 없음.", "", "8.1.2 공급망 추적성 검토"])
        self.assertEqual(E.drop_blank_after_note(document), 1)
        self.assertEqual([p.text for p in document.paragraphs],
                         ["* [RSF102] 플루오로메톨론 사용이력 없음.", "8.1.2 공급망 추적성 검토"])

    def test_보통_글_뒤의_빈_줄은_둔다(self):
        document = self._doc(["제품품질평가는 2025년도 …", "", "5. 책임과 권한"])
        self.assertEqual(E.drop_blank_after_note(document), 0)


class 항_사이_한_줄(unittest.TestCase):
    """담당자 지시(2026-09): "16항 아래와 17항 시작부분은 한칸 띄워줘야 돼. 모든 항이 마찬가지야." """

    def _doc(self, lines):
        document = docx.Document()
        for text in lines:
            document.add_paragraph(text)
        return document

    def test_항_제목_앞에_빈_줄을_넣는다(self):
        document = self._doc(["16. 결론", "…확인하였음.", "17. 참고 자료"])
        self.assertEqual(E.space_before_sections(document), 1)
        self.assertEqual([p.text for p in document.paragraphs],
                         ["16. 결론", "…확인하였음.", "", "17. 참고 자료"])

    def test_이미_빈_줄이_있으면_더_넣지_않는다(self):
        document = self._doc(["…확인하였음.", "", "17. 참고 자료"])
        self.assertEqual(E.space_before_sections(document), 0)

    def test_작은_제목은_그대로_둔다(self):
        document = self._doc(["* [RSF102] 사용이력 없음.", "8.1.2 공급망 추적성 검토"])
        self.assertEqual(E.space_before_sections(document), 0)


class 변경사항_간추리기(unittest.TestCase):
    def test_큰_항목만_남고_잔가지는_버린다(self):
        got = brief_change(
            "1. 안연고 튜브 자재 시험 검체 수량 확대\n"
            ": 시험 검체 수량을 125EA 에서 315EA 로 확대하고자 함\n"
            '1) "QC-142 자재 검체 채취 규정" 개정\n'
            "- 성상, 시각적 검사 항목 시험수량 변경\n"
            "2. 튜브 막힘을 확인하는 시험 항목 추가\n"
            ": 명확한 기술이 없어\n"
            '"튜브개봉시험"을 추가 하고자 함\n'
            '1) "QC-201 중간 공정 검사 규정" 개정')
        self.assertEqual(got, [
            "1. 안연고 튜브 자재 시험 검체 수량 확대: 시험 검체 수량을 125EA 에서 315EA 로 확대하고자 함",
            '2. 튜브 막힘을 확인하는 시험 항목 추가: 명확한 기술이 없어 "튜브개봉시험"을 추가 하고자 함'])

    def test_항목_번호가_없으면_원문_그대로(self):
        self.assertEqual(brief_change("MAINMIXER 및 PREMIXER SHT 신규 설정"),
                         ["MAINMIXER 및 PREMIXER SHT 신규 설정"])


if __name__ == "__main__":
    unittest.main()


class 변경관리_해당_여부(unittest.TestCase):
    """담당자가 12항에 넣은 변경관리는 모두 그 제품에 해당한다 (담당자 2026-09).

    '관련 제품' 은 제품 이름을 하나하나 적기보다 무리로 적는 일이 잦다.
    """
    성분 = ["플루오로메톨론", "겐타마이신황산염"]

    def _본다(self, products, title=""):
        from pqr.engine.recipe_ointment import change_covers
        return change_covers({"products": products, "title": title}, "디겐타안연고", self.성분)

    def test_무리로_적은_것은_해당한다(self):
        for 관련 in ("안연고 전 제품", "안연고제 전제품 일괄", "무균라인에서 생산하는 전 제품"):
            self.assertTrue(self._본다(관련), 관련)

    def test_주성분을_건드리는_변경은_해당한다(self):
        """대상이 후메론점안액으로만 적혀 있어도 플루오로메톨론 변경은 디겐타에도 온다."""
        self.assertTrue(self._본다("후메론점안액 (다회용, 일회용)", "플루오로메톨론 멸균 온도 변경 건"))

    def test_제품_이름이_적혀_있으면_해당한다(self):
        self.assertTrue(self._본다("디겐타안연고, 퀴노비드안연고"))

    def test_남의_제품만_적힌_것은_문의로_남긴다(self):
        self.assertFalse(self._본다("메섹신정", "정제 코팅 조건 변경"))
