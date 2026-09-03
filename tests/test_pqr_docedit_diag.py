# -*- coding: utf-8 -*-
"""표 칸의 사선(↗) 판별 — Word 는 '사선 없음'도 XML 에 적어 둔다."""
from __future__ import unicode_literals

import unittest

import docx
from docx.oxml.ns import qn

from pqr.engine import docedit as E


def _cell():
    """칸 하나짜리 표를 만들어 그 칸을 돌려준다."""
    doc = docx.Document()
    return doc, doc.add_table(rows=1, cols=1).cell(0, 0)


def _border(cell, tag, val):
    pr = E._tcpr(cell._tc)
    borders = E.get_or_add(pr, "tcBorders")
    el = E.get_or_add(borders, tag)
    if val is not None:
        el.set(qn("w:val"), val)
    return el


class 사선판별(unittest.TestCase):
    def test_테두리가_없으면_사선도_없다(self):
        _, cell = _cell()
        self.assertFalse(E._has_diag(cell._tc))

    def test_val_nil_은_사선이_아니다(self):
        """Word 가 '사선 없음'으로 적어 두는 형태 — 요소만 보고 있다고 하면 안 된다."""
        for val in ("nil", "none"):
            _, cell = _cell()
            _border(cell, "tr2bl", val)
            self.assertFalse(E._has_diag(cell._tc), val)

    def test_val_single_은_사선이다(self):
        for tag in ("tr2bl", "tl2br"):
            _, cell = _cell()
            _border(cell, tag, "single")
            self.assertTrue(E._has_diag(cell._tc), tag)

    def test_val_이_없으면_사선으로_본다(self):
        _, cell = _cell()
        _border(cell, "tr2bl", None)
        self.assertTrue(E._has_diag(cell._tc))

    def test_add_diag_로_그은_사선을_다시_읽는다(self):
        _, cell = _cell()
        E.add_diag(cell)
        self.assertTrue(E._has_diag(cell._tc))

    def test_nil_인_빈_칸에는_사선을_새로_긋는다(self):
        """val=nil 을 사선으로 잘못 읽으면 이 칸이 빈 채로 남는다."""
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "값"
        _border(table.cell(0, 1), "tr2bl", "nil")
        self.assertEqual(E.diag_all_empty(doc), 1)
        self.assertTrue(E._has_diag(table.cell(0, 1)._tc))
        self.assertFalse(E._has_diag(table.cell(0, 0)._tc))


if __name__ == "__main__":
    unittest.main()
