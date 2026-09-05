# -*- coding: utf-8 -*-
"""16항 진행 중 일탈 문단 · 17항 Cpk 경향분석 Sheet 참고 줄 (2026-09 점검, 한림 결재본 PQR25 퀴노비드안연고)."""
from __future__ import unicode_literals

import unittest

import docx

from pqr.engine import conclusion
from pqr.engine import recipe_ointment as R


class 진행중_일탈(unittest.TestCase):
    def test_진행_중인_것이_없으면_문단_없음(self):
        self.assertIsNone(conclusion.deviation_sentence(["DR-1"], []))

    def test_결재본_문안대로(self):
        got = conclusion.deviation_sentence(["DR-240828-03"], ["DR-240827-01", "DR-241202-03"])
        self.assertEqual(got, "점검기간 동안 발생한 내수용 제품의 일탈 3건 중 1건(DR-240828-03)은 적절하게 조치 완료됨을 "
                              "확인하였으며 현재 조치 사항 진행 중인 2건(DR-240827-01, DR-241202-03)에 대해서는 "
                              "차년도 제품품질평가에서 조치 사항이 적절하게 완료되었는지 확인하도록 한다.")

    def test_Cpk_문제_없고_진행_중_일탈이_있으면_16_1_16_2(self):
        got = conclusion.texts("디겐타안연고", "디겐타", True, 12, 2025, 2026, 4, [], "진행 중 문단")
        self.assertEqual(len(got), 2)
        self.assertTrue(got[0].startswith("16.1 디겐타안연고에 대한"))
        self.assertEqual(got[1], "16.2 진행 중 문단")

    def test_둘_다_없으면_번호_없는_한_문단(self):
        got = conclusion.texts("디겐타안연고", "디겐타", True, 3, 2025, 2026, 4, [], None)
        self.assertEqual(len(got), 1)
        self.assertFalse(got[0].startswith("16."))


def doc_with_refs(*lines):
    d = docx.Document()
    d.add_paragraph("17. 참고 자료")
    for line in lines:
        d.add_paragraph(line)
    d.add_paragraph("18. 첨부 문서")
    return d


BASE = ("- 제품표준서 MF-7014", "- QC-126 제품품질평가규정",
        "- 의약품 제조 및 품질관리에 관한 규정 고시 별표 17(완제의약품 제조)")


class Cpk_참고줄(unittest.TestCase):
    def refs(self, d):
        out, on = [], False
        for p in d.paragraphs:
            if p.text.startswith("17."):
                on = True; continue
            if p.text.startswith("18."):
                break
            if on:
                out.append(p.text)
        return out

    def test_10_Lot_이상이면_별표_17_뒤에_두_줄(self):
        d = doc_with_refs(*BASE)
        ref = next(p for p in d.paragraphs if p.text.startswith("17."))
        R._cpk_references(d, ref, True, lambda m: None)
        got = self.refs(d)
        self.assertEqual(got[-2:], list(R.CPK_SHEETS))
        self.assertEqual(got[2], BASE[2])

    def test_이미_있으면_더_넣지_않는다(self):
        d = doc_with_refs(*(BASE + R.CPK_SHEETS))
        ref = next(p for p in d.paragraphs if p.text.startswith("17."))
        R._cpk_references(d, ref, True, lambda m: None)
        self.assertEqual(len(self.refs(d)), 5)

    def test_10_Lot_미만이면_뺀다(self):
        d = doc_with_refs(*(BASE + R.CPK_SHEETS))
        ref = next(p for p in d.paragraphs if p.text.startswith("17."))
        R._cpk_references(d, ref, False, lambda m: None)
        self.assertEqual(self.refs(d), list(BASE))


if __name__ == "__main__":
    unittest.main()
