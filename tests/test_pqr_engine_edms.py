# -*- coding: utf-8 -*-
"""EDMS 결재본 서식(E-HLF-32) 찾기와 빈 쪽 방지 정리 — 합성 문서로만 시험한다."""
from __future__ import unicode_literals

import os
import shutil
import tempfile
import unittest

import docx
from docx.oxml.ns import qn

from pqr.engine import docedit as E
from pqr.engine import edms, layout


def make_docx(path, footer="", body=("본문",)):
    d = docx.Document()
    for line in body:
        d.add_paragraph(line)
    if footer:
        d.sections[0].footer.paragraphs[0].text = footer
    d.save(path)
    return path


class 서식알아보기(unittest.TestCase):
    def setUp(self):
        self.dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.dir, ignore_errors=True)

    def test_바닥글에_EHLF32_가_있으면_서식이다(self):
        p = make_docx(os.path.join(self.dir, "a.docx"), footer="EHLF-32/Rev.000")
        self.assertTrue(edms.is_edms_form(p))

    def test_전년도_결재본_바닥글은_서식이_아니다(self):
        p = make_docx(os.path.join(self.dir, "b.docx"), footer="한림제약 HLF-QC-126-01/Rev.010-1")
        self.assertFalse(edms.is_edms_form(p))

    def test_docx_가_아니면_서식이_아니다(self):
        p = os.path.join(self.dir, "c.doc")
        open(p, "wb").write(b"\xd0\xcf\x11\xe0")
        self.assertFalse(edms.is_edms_form(p))
        self.assertFalse(edms.is_edms_form(None))

    def test_layout_도_같은_기준으로_알아본다(self):
        p = make_docx(os.path.join(self.dir, "a.docx"), footer="EHLF-32/Rev.000")
        self.assertTrue(layout.is_edms(docx.Document(p)))
        q = make_docx(os.path.join(self.dir, "b.docx"), footer="HLF-QC-126-01")
        self.assertFalse(layout.is_edms(docx.Document(q)))


class 서식찾기(unittest.TestCase):
    def setUp(self):
        self.root = tempfile.mkdtemp()
        self.product = os.path.join(self.root, "QC1-7007 퀴노비드안연고")
        os.makedirs(self.product)

    def tearDown(self):
        shutil.rmtree(self.root, ignore_errors=True)

    def test_제품_폴더의_서식을_찾는다(self):
        p = make_docx(os.path.join(self.product, "EDMS 결재본.docx"), footer="EHLF-32/Rev.000")
        make_docx(os.path.join(self.product, "16. 전년도 PQR.docx"), footer="HLF-QC-126-01")
        self.assertEqual(edms.find_form(self.product), p)

    def test_하위_폴더와_공통_폴더도_본다(self):
        sub = os.path.join(self.product, "16. 전년도 PQR")
        os.makedirs(sub)
        p = make_docx(os.path.join(sub, "서식.docx"), footer="EHLF-32/Rev.000")
        self.assertEqual(edms.find_form(self.product), p)
        os.remove(p)
        common = os.path.join(self.root, "공통")
        os.makedirs(common)
        q = make_docx(os.path.join(common, "E-HLF-32 서식 Rev.000.docx"), footer="EHLF-32/Rev.000")
        self.assertEqual(edms.find_form(self.product), q)

    def test_완성본과_제출본은_서식으로_보지_않는다(self):
        make_docx(os.path.join(self.product, "QC1-7007 완성본.docx"), footer="EHLF-32/Rev.000")
        make_docx(os.path.join(self.product, "제출용 보고서.docx"), footer="EHLF-32/Rev.000")
        got = edms.find_form(self.product)
        self.assertTrue(edms.is_shipped(got), got)       # 결과물은 건너뛰고 프로그램 껍데기로

    def test_폴더에_없으면_프로그램에_든_껍데기를_쓴다(self):
        got = edms.find_form(self.product)
        self.assertTrue(edms.is_shipped(got), got)
        self.assertTrue(edms.is_edms_form(got))
        self.assertIsNone(edms.find_form(os.path.join(self.root, "없는폴더")))

    def test_프로그램에_든_껍데기는_본문이_없다(self):
        d = docx.Document(edms.SHIPPED_FORM)
        texts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        self.assertEqual(texts[0], "목차 (Table of Contents)")
        self.assertTrue(texts[-1].startswith("1."))
        self.assertEqual(len(texts), 2)                      # 제품 고유 글이 없다
        self.assertEqual(len(d.tables), 1)                   # 목차 표뿐

    def test_바탕_고르기(self):
        self.assertEqual(edms.choose_base("form.docx", "prev.doc")[0], "form.docx")
        self.assertEqual(edms.choose_base(None, "prev.doc")[0], "prev.doc")
        base, why = edms.choose_base(None, None)
        self.assertIsNone(base)
        self.assertIn("E-HLF-32", why)


def _blank(doc):
    return doc.add_paragraph("")


def _hard_break(doc):
    p = doc.add_paragraph("")
    p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    return p


class 빈쪽방지(unittest.TestCase):
    def test_쪽나눔_앞_빈_문단을_지우고_나눔을_옮긴다(self):
        d = docx.Document()
        d.add_paragraph("6. 제조내역")
        _blank(d); _blank(d)
        _hard_break(d)
        d.add_paragraph("7. 수율현황표")
        got = E.tidy_page_breaks(d)
        texts = [p.text for p in d.paragraphs]
        self.assertEqual(texts, ["6. 제조내역", "7. 수율현황표"])
        self.assertEqual(got["앞 빈 문단"], 2)
        self.assertEqual(got["쪽나눔 전환"], 1)
        pr = d.paragraphs[1]._p.find(qn("w:pPr"))
        self.assertIsNotNone(pr.find(qn("w:pageBreakBefore")))

    def test_표_앞_쪽나눔은_그대로_둔다(self):
        d = docx.Document()
        d.add_paragraph("9.2.1 조제 완료 후")
        _blank(d)
        _hard_break(d)
        d.add_table(rows=1, cols=1)
        got = E.tidy_page_breaks(d)
        self.assertEqual(got["쪽나눔 전환"], 0)
        self.assertEqual(got["앞 빈 문단"], 1)
        self.assertEqual(len(d.paragraphs), 2)      # 제목 + 쪽나눔 문단
        self.assertTrue(E._has_page_break(d.paragraphs[1]._p))

    def test_문서_끝_빈_문단을_지운다(self):
        d = docx.Document()
        d.add_paragraph("18. 첨부 문서")
        _blank(d); _blank(d)
        got = E.tidy_page_breaks(d)
        self.assertEqual(got["끝 빈 문단"], 2)
        self.assertEqual([p.text for p in d.paragraphs], ["18. 첨부 문서"])

    def test_표_바로_뒤의_빈_문단은_남긴다(self):
        d = docx.Document()
        d.add_paragraph("15. 시정조치사항")
        d.add_table(rows=1, cols=1)
        _blank(d)
        got = E.tidy_page_breaks(d)
        self.assertEqual(got["끝 빈 문단"], 0)
        self.assertEqual(len(d.paragraphs), 2)


class 머리글라벨(unittest.TestCase):
    """EDMS 머리글(문서번호·Rev. No.·Page)에서 쪽 번호 칸에 날짜를 덮어쓰지 않아야 한다."""

    def test_edms_머리글은_Page_칸을_건드리지_않는다(self):
        import datetime as dt
        import re
        d = docx.Document()
        hdr = d.sections[0].header
        t = hdr.add_table(rows=3, cols=2, width=docx.shared.Inches(6))
        for i, label in enumerate(("문서번호", "Rev. No.", "Page")):
            t.cell(i, 0).text = label
        t.cell(2, 1).text = "3 / 32"
        # recipe 의 머리글 처리와 같은 논리
        today = dt.date(2026, 9, 3)
        write_year = today.year
        new_no, full_name = "", "짧은이름"
        for row in t.rows:
            cells = E.raw_cells(row)
            labels = "".join(E.cell_text(c) for c in cells[:-1])
            target = cells[-1]
            if "문서번호" in labels:
                old_no = E.cell_text(target).strip()
                if re.search(r"PQR\d{2}-", old_no):
                    new_no = re.sub(r"PQR\d{2}-", "PQR%02d-" % (write_year % 100), old_no)
                    E.set_cell(target, new_no)
            elif "작성일자" in labels:
                E.set_cell(target, today.strftime("%Y.%m.%d"))
            elif "제품명" in labels:
                full_name = E.cell_text(target).strip() or full_name
        self.assertEqual(E.cell_text(t.cell(2, 1)).strip(), "3 / 32")
        self.assertEqual(E.cell_text(t.cell(0, 1)).strip(), "")
        self.assertEqual(new_no, "")
        self.assertEqual(full_name, "짧은이름")



class 칸폭(unittest.TestCase):
    def _table(self, widths, rows):
        d = docx.Document()
        t = d.add_table(rows=len(rows), cols=len(widths))
        grid = t._tbl.find(qn("w:tblGrid"))
        for g, w in zip(grid.findall(qn("w:gridCol")), widths):
            g.set(qn("w:w"), str(w))
        for ri, row in enumerate(rows):
            for ci, txt in enumerate(row):
                t.cell(ri, ci).text = txt
        return t

    def _grid(self, t):
        return [int(g.get(qn("w:w"))) for g in t._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol"))]

    def test_고정_배치와_그리드_폭(self):
        t = self._table([800, 1200, 4000, 4000], [["연번", "Lot No.", "성상", "생균수"], ["1", "A", "b", "c"]])
        n = E.fix_table_widths(t)
        self.assertEqual(n, 8)
        pr = t._tbl.find(qn("w:tblPr"))
        self.assertEqual(pr.find(qn("w:tblLayout")).get(qn("w:type")), "fixed")
        w = [c._tc.find(qn("w:tcPr")).find(qn("w:tcW")) for c in E.raw_cells(t.rows[1])]
        self.assertEqual([x.get(qn("w:w")) for x in w], ["800", "1200", "4000", "4000"])
        self.assertEqual({x.get(qn("w:type")) for x in w}, {"dxa"})

    def test_한_열만_길면_다시_나눈다(self):
        long = "함량시험의 검액과 표준액을 Diode array detector 로 200~400nm 에서 측정한 결과 주피크 UV spectrum 은 동일하다."
        t = self._table([800, 1200, 2400, 2300, 3100],
                        [["연번", "Lot No.", "성상", "1)", "2)"],
                         ["1", "OEY101", "무취의 담황색 안연고제", "검액은 표준액과 동일한 주피크 유지시간을 나타낸다.", long]])
        new = E.balance_columns(t)
        self.assertIsNotNone(new)
        self.assertEqual(new[:2], [800, 1200])           # 열쇠 열은 그대로
        self.assertEqual(sum(new), 800 + 1200 + 2400 + 2300 + 3100)
        self.assertGreater(new[4], new[2])                # 긴 열이 넓어진다
        self.assertGreaterEqual(new[2], int(0.6 * 2400))  # 60% 바닥

    def test_이미_한_줄이면_손대지_않는다(self):
        t = self._table([800, 1200, 3000, 3000],
                        [["연번", "Lot No.", "함량(%)", "입자도\n(㎛ 이하)"], ["1", "OEY101", "107.3", "61"]])
        self.assertIsNone(E.balance_columns(t))
        self.assertEqual(self._grid(t), [800, 1200, 3000, 3000])

    def test_머리행의_긴_제목은_세지_않는다(self):
        t = self._table([800, 1200, 1500, 1500],
                        [["연번", "Lot No.", "50㎛ 이상 이물 총합계(개)", "개개 페트리접시 중 8개 초과(매)"],
                         ["1", "OEY101", "0", "0"]])
        self.assertIsNone(E.balance_columns(t))

    def test_글_길이(self):
        self.assertEqual(E.text_units("abc"), 3)
        self.assertEqual(E.text_units("한글"), 4)
        self.assertEqual(E.text_units("한 a\n한글한글"), 8)

if __name__ == "__main__":
    unittest.main()


class 서식에옮겨담기(unittest.TestCase):
    """채운 문서의 본문을 EDMS 서식 껍데기(머리글·바닥글·목차·쪽 설정)에 담는다."""

    def setUp(self):
        self.dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.dir, ignore_errors=True)

    def _filled(self):
        d = docx.Document()
        d.add_paragraph("퀴노비드안연고 제품품질평가 보고서")          # 표지
        d.add_paragraph("검토 및 승인")
        t = d.add_table(rows=2, cols=3); t.cell(0, 0).text = "구 분"; t.cell(0, 1).text = "서 명"; t.cell(0, 2).text = "서명일자"
        d.add_paragraph("개 정 내 역")
        d.add_paragraph("1. 목 적")
        d.add_paragraph("당사 의약품의 품질에 대하여 …")
        d.add_paragraph("6. 제조내역 확인")
        t6 = d.add_table(rows=2, cols=2); t6.cell(0, 0).text = "연번"; t6.cell(0, 1).text = "Lot No."; t6.cell(1, 1).text = "OEY101"
        d.sections[0].footer.paragraphs[0].text = "한림제약 HLF-QC-126-01/Rev.010-1"
        d.sections[0].header.paragraphs[0].text = "PQR26-2-QUIO3"
        p = os.path.join(self.dir, "filled.docx"); d.save(p); return p

    def _form(self):
        d = docx.Document()
        d.add_paragraph("목차 (Table of Contents)")
        t = d.add_table(rows=2, cols=2); t.cell(0, 0).text = "1.  목 적 (Purpose)"; t.cell(0, 1).text = "2"
        t.cell(1, 0).text = "6.  제조내역 확인"; t.cell(1, 1).text = "4"
        d.add_paragraph(""); d.add_paragraph("")
        d.add_paragraph("1. 목 적")
        d.add_paragraph("서식의 목적 글")
        d.sections[0].footer.paragraphs[0].text = "EHLF-32/Rev.000"
        d.sections[0].header.paragraphs[0].text = "제품 품질 평가 보고서 문서번호 Rev. No. Page"
        d.sections[0].left_margin = docx.shared.Mm(19)
        p = os.path.join(self.dir, "form.docx"); d.save(p); return p

    def test_본문은_남기고_껍데기만_바꾼다(self):
        from pqr.engine import rehouse
        out = os.path.join(self.dir, "out.docx")
        got = rehouse.rehouse(self._filled(), self._form(), out)
        self.assertEqual(got["앞부분 삭제"], 4)          # 표지·결재표 제목·결재표·개정 내역
        self.assertGreaterEqual(got["머리글·바닥글"], 2)
        d = docx.Document(out)
        texts = [p.text for p in d.paragraphs if p.text.strip()]
        self.assertEqual(texts[0], "목차 (Table of Contents)")
        self.assertIn("1. 목 적", texts)                  # 채운 문서의 1항이 남는다
        self.assertNotIn("서식의 목적 글", texts)          # 서식의 본문은 안 들어온다
        self.assertNotIn("검토 및 승인", texts)
        self.assertEqual([c.text for c in d.tables[0].rows[0].cells], ["1.  목 적 (Purpose)", "2"])
        self.assertEqual(d.tables[1].cell(1, 1).text, "OEY101")
        self.assertIn("EHLF-32", d.sections[0].footer.paragraphs[0].text)
        self.assertIn("문서번호", d.sections[0].header.paragraphs[0].text)
        self.assertLess(abs(d.sections[0].left_margin - docx.shared.Mm(19)), 1000)   # twip 반올림
        self.assertTrue(layout.is_edms(d))
        # 1항 제목은 '앞에서 쪽 나눔' 으로 목차 다음 쪽에서 시작한다
        head = next(p for p in d.paragraphs if p.text.startswith("1. 목 적"))
        self.assertIsNotNone(head._p.find(qn("w:pPr")).find(qn("w:pageBreakBefore")))

    def test_1항_제목이_없으면_실패한다(self):
        from pqr.engine import rehouse
        d = docx.Document(); d.add_paragraph("아무 제목 없음")
        bad = os.path.join(self.dir, "bad.docx"); d.save(bad)
        with self.assertRaises(rehouse.RehouseError):
            rehouse.rehouse(bad, self._form(), os.path.join(self.dir, "o.docx"))


class 압축속결재본(unittest.TestCase):
    def setUp(self):
        self.dir = tempfile.mkdtemp()
        self.product = os.path.join(self.dir, "QC1-7014 디겐타안연고")
        os.makedirs(self.product)

    def tearDown(self):
        shutil.rmtree(self.dir, ignore_errors=True)

    def test_zip_안의_doc_를_꺼낸다(self):
        import zipfile
        from pqr.engine import writer
        z = os.path.join(self.product, "16 전년도 PQR word & excel (PQR25).zip")
        with zipfile.ZipFile(z, "w") as zf:
            zf.writestr("a. PQR25 디겐타안연고.doc", b"\xd0\xcf\x11\xe0 fake")
            zf.writestr("a. 함량 Cpk 계산 파일.xls", b"x")
        work = tempfile.mkdtemp()
        got = writer.find_previous(self.product, work)
        self.assertTrue(got and got.endswith(".doc"))
        self.assertTrue(os.path.isfile(got))
        # Cpk 엑셀은 꺼낸 결재본과 같은 폴더(압축을 푼 곳)에서 찾는다
        src = writer.attachment_source(self.product, work)
        self.assertTrue(src.endswith(".zip") or any(
            n.lower().endswith(".xls") for n in os.listdir(os.path.dirname(src))))

    def test_자료_전체를_묶은_zip_하나만_있어도_찾는다(self):
        """디겐타안연고: '2026년 PQR 필요 자료.zip' 하나에 16. 전년도 PQR 압축까지 들어 있다."""
        import zipfile, io as _io
        from pqr.engine import writer
        inner = _io.BytesIO()
        with zipfile.ZipFile(inner, "w") as zf:
            zf.writestr("a. PQR25 디겐타안연고.doc", b"\xd0\xcf\x11\xe0 fake")
            zf.writestr("a. 함량 Cpk 계산 파일.xls", b"x")
        bundle = os.path.join(self.product, "디겐타 안연고 2026년 PQR 필요 자료.zip")
        with zipfile.ZipFile(bundle, "w") as zf:
            zf.writestr("디겐타 안연고 2026년 PQR 필요 자료/3. 허가증.pdf", b"%PDF")
            zf.writestr("디겐타 안연고 2026년 PQR 필요 자료/16 전년도 PQR word & excel (PQR25).zip", inner.getvalue())
        work = tempfile.mkdtemp()
        got = writer.find_previous(self.product, work)
        self.assertTrue(got and got.endswith("a. PQR25 디겐타안연고.doc"), got)

    def test_docx_가_있으면_압축보다_먼저(self):
        import zipfile
        from pqr.engine import writer
        make_docx(os.path.join(self.product, "16. 전년도 PQR.docx"), footer="HLF-QC-126-01")
        with zipfile.ZipFile(os.path.join(self.product, "16 PQR25.zip"), "w") as zf:
            zf.writestr("a.doc", b"x")
        got = writer.find_previous(self.product, tempfile.mkdtemp())
        self.assertTrue(got.endswith("16. 전년도 PQR.docx"))

    def test_없으면_None(self):
        from pqr.engine import writer
        self.assertIsNone(writer.find_previous(self.product, tempfile.mkdtemp()))
        self.assertIsNone(writer.attachment_source(self.product))
