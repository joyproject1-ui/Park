# -*- coding: utf-8 -*-
"""16항 결론 — 배포본 'PQR 작성방법 공유의 건'(2026-09-04) 문안."""
from __future__ import unicode_literals

import unittest

import docx

from pqr.engine import conclusion as C


class 문안(unittest.TestCase):
    def test_Cpk_1_미만이_없으면_번호_없는_종합_문장_하나(self):
        # 문단이 하나뿐이면 번호를 붙이지 않는다 (담당자 확인 2026-09, 2026 디겐타 결재본).
        got = C.texts("퀴노비드안연고(오플록사신)", "퀴노비드안연고", True, 17, 2025, 2026, 4, [])
        self.assertEqual(len(got), 1)
        self.assertTrue(got[0].startswith("퀴노비드안연고(오플록사신)에 대한 제품품질평가 결과"))
        self.assertIn("해당되는 사항은 없음을 확인하였음.", got[0])
        self.assertNotIn("10 Lot", got[0])                  # 미산출 문장은 결론에 쓰지 않는다
        self.assertNotIn("Cpk", got[0])

    def test_Cpk_1_미만이_있으면_네_문단(self):
        got = C.texts("퀴노비드안연고(오플록사신)", "퀴노비드안연고", True, 17, 2025, 2026, 4,
                      [("함량", 0.93), ("입자도", 0.79)])
        self.assertEqual(len(got), 4)
        self.assertTrue(got[0].startswith("16.1 본 제품 품질 평가를 통해 공정능력을 평가한 결과 Cpk 값이 1 미만인 항목이 아래 표와 같이"))
        self.assertEqual(got[1], "16.2 2025년 생산된 17Lot에 사용된 원료 검토 결과, 모든 항목에서 기준 내 적합한 결과였으며 제품 시험 검토 결과 허가 및 자가 기준 내 적합함을 확인함.")
        self.assertIn("2026년도 4분기 내 ‘공정능력지수 검토 계획서(HLF-QC-126-10)’을 작성하여, 퀴노비드안연고 공정 개선", got[2])
        self.assertTrue(got[3].startswith("16.4 퀴노비드안연고(오플록사신)에 대한 제품품질평가 결과"))

    def test_생산이력_없음(self):
        got = C.texts("아이퓨어점안액", "아이퓨어점안액", False, 0, 2025, 2026, 4, [])
        self.assertEqual(len(got), 1)
        self.assertIn("2025년 생산 이력이 없어", got[0])
        self.assertIn("추후 생산 시", got[0])

    def test_1_미만_항목만_표에_오른다(self):
        rows = C.low_cpk_items({"assay": 0.93, "particle": 0.79, "metal": 24.15})
        self.assertEqual(rows, [("함량", 0.93), ("입자도", 0.79)])
        self.assertEqual(C.low_cpk_items({"assay": 1.13, "particle": None}), [])


class 문서에넣기(unittest.TestCase):
    def _doc(self):
        d = docx.Document()
        d.add_paragraph("15. 시정조치사항")
        d.add_paragraph("16. 결론")
        d.add_paragraph("16.1. 옛 결론 문장")
        d.add_paragraph("16.2. 옛 일탈 문장")
        d.add_paragraph("17. 참고 자료")
        d.add_paragraph("- 제품표준서")
        return d

    def test_옛_16x_문단을_지우고_새_문안을_넣는다(self):
        d = self._doc()
        n = C.apply(d, "퀴노비드안연고(오플록사신)", "퀴노비드안연고", True, 17, 2025, 2026, 4, {"assay": 1.2})
        self.assertEqual(n, 1)
        texts = [p.text for p in d.paragraphs]
        self.assertEqual(texts[1], "16. 결론")
        self.assertTrue(texts[2].startswith("퀴노비드안연고(오플록사신)에 대한"))
        self.assertEqual(texts[3], "17. 참고 자료")
        self.assertEqual(len(d.tables), 0)

    def test_Cpk_1_미만이면_표가_16_1_아래에_생긴다(self):
        d = self._doc()
        n = C.apply(d, "퀴노비드안연고(오플록사신)", "퀴노비드안연고", True, 17, 2025, 2026, 4,
                    {"assay": 0.93, "particle": 0.79, "metal": 24.15})
        self.assertEqual(n, 4)
        self.assertEqual(len(d.tables), 1)
        t = d.tables[0]
        self.assertEqual([c.text for c in t.rows[0].cells], ["항목", "판정 기준", "Cpk"])
        self.assertEqual(t.cell(1, 0).text, "함량"); self.assertEqual(t.cell(1, 2).text, "0.93")
        self.assertEqual(t.cell(2, 0).text, "입자도"); self.assertEqual(t.cell(2, 2).text, "0.79")
        self.assertIn("공정능력 충분", t.cell(1, 1).text)
        # 본문 순서: 16. 결론 → 16.1 → [표] → 16.2 → 16.3 → 16.4 → 17.
        body = d.element.body
        kinds = []
        for el in body:
            tag = el.tag.split("}")[1]
            if tag == "tbl":
                kinds.append("표")
            elif tag == "p":
                t_ = "".join(x.text or "" for x in el.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
                if t_.strip():
                    kinds.append(t_[:4])
        i = kinds.index("16. ")
        self.assertEqual(kinds[i:i + 7], ["16. ", "16.1", "표", "16.2", "16.3", "16.4", "17. "])

    def test_제목이_없으면_0(self):
        d = docx.Document(); d.add_paragraph("아무것도")
        self.assertEqual(C.apply(d, "a", "a", True, 1, 2025, 2026, 4, {}), 0)


if __name__ == "__main__":
    unittest.main()


class 계획서분기(unittest.TestCase):
    def test_작성일의_다음_분기(self):
        import datetime as dt
        self.assertEqual(C.plan_quarter(dt.date(2026, 9, 4)), (2026, 4))
        self.assertEqual(C.plan_quarter(dt.date(2026, 1, 15)), (2026, 2))
        self.assertEqual(C.plan_quarter(dt.date(2026, 11, 2)), (2027, 1))
