# -*- coding: utf-8 -*-
"""7항 비고 열 — 2025 결재본처럼 자료 줄 한 묶음·요약 줄 한 묶음에 사선 하나씩 (담당자 2026-09)."""
from __future__ import unicode_literals

import unittest

import docx
from docx.oxml.ns import qn

from pqr.engine import docedit as E


class _Raw(object):
    """python-docx 의 cell(r, c) 는 세로 병합된 칸을 위 칸으로 돌려준다 — 원래 tc 를 본다."""
    def __init__(self, tc):
        self._tc = tc


def raw(t, r, c):
    return _Raw(t._tbl.findall(qn("w:tr"))[r].findall(qn("w:tc"))[c])


def vmerge(cell):
    pr = cell._tc.find(qn("w:tcPr"))
    vm = pr.find(qn("w:vMerge")) if pr is not None else None
    if vm is None:
        return None
    return vm.get(qn("w:val")) or "continue"


def make():
    d = docx.Document()
    t = d.add_table(rows=9, cols=6)
    heads = ("연번", "공정", "조제", "충전", "포장", "비고")
    for c, h in enumerate(heads):
        t.cell(0, c).text = h
    t.cell(1, 1).text = "공정"; t.cell(2, 1).text = "기준 Lot No."
    for r, lot in ((3, "OGY301"), (4, "OGY901"), (5, "OGY902")):
        t.cell(r, 0).text = str(r - 2); t.cell(r, 1).text = lot; t.cell(r, 2).text = "96.89"
    for r, h in ((6, "최댓값"), (7, "최솟값"), (8, "평균")):
        t.cell(r, 0).text = h; t.cell(r, 2).text = "99.21"
    # 머리글 비고는 세 줄 세로 병합
    E.set_vmerge(t.cell(0, 5), "restart"); E.set_vmerge(t.cell(1, 5), None); E.set_vmerge(t.cell(2, 5), None)
    return d, t


class 비고_두_묶음(unittest.TestCase):
    def test_자료_줄과_요약_줄이_따로_합쳐지고_사선은_둘(self):
        _, t = make()
        self.assertEqual(E.merge_remark_column(t), 6)
        self.assertEqual([vmerge(raw(t, r, 5)) for r in range(9)],
                         ["restart", "continue", "continue",          # 머리글 그대로
                          "restart", "continue", "continue",          # 자료 줄 묶음
                          "restart", "continue", "continue"])         # 요약 줄 묶음
        self.assertTrue(E.has_diag(raw(t, 3, 5)))
        self.assertTrue(E.has_diag(raw(t, 6, 5)))
        self.assertFalse(E.has_diag(raw(t, 1, 5)))
        self.assertFalse(E.has_diag(raw(t, 4, 5)))

    def test_예전_코드가_망친_머리글도_되돌린다(self):
        _, t = make()
        E.set_vmerge(raw(t, 1, 5), "restart"); E.add_diag(raw(t, 1, 5))
        E.merge_remark_column(t)
        self.assertEqual(vmerge(raw(t, 1, 5)), "continue")
        self.assertFalse(E.has_diag(raw(t, 1, 5)))

    def test_자료가_한_줄이면_병합_없이_사선만(self):
        d = docx.Document()
        t = d.add_table(rows=5, cols=3)
        for c, h in enumerate(("연번", "조제", "비고")):
            t.cell(0, c).text = h
        t.cell(1, 0).text = "1"; t.cell(1, 1).text = "98.78"
        for r, h in ((2, "최댓값"), (3, "최솟값"), (4, "평균")):
            t.cell(r, 0).text = h
        self.assertEqual(E.merge_remark_column(t), 4)
        self.assertIsNone(vmerge(raw(t, 1, 2)))
        self.assertTrue(E.has_diag(raw(t, 1, 2)))
        self.assertEqual(vmerge(raw(t, 2, 2)), "restart")


if __name__ == "__main__":
    unittest.main()
