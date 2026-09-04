# -*- coding: utf-8 -*-
"""표 차림새 — 2026 디겐타 결재본을 보고 담당자가 짚어 준 것들.

  · 폴더와 압축에 같은 파일이 있으면 한 번만 읽는다 (8.2 표에 같은 줄이 두 벌 실렸다)
  · 같은 원료·코드는 시험번호가 달라도 세로로 합친다
  · 표는 모두 '창에 자동으로 맞춤' — 본문 폭에 딱 맞춘다
  · 내용이 없으면 빈 줄을 하나만 남긴다
"""
import os
import shutil
import tempfile
import unittest
import zipfile

import docx
from docx.oxml.ns import qn

from pqr.engine import collect, docedit as E


class 같은_파일은_한_번만(unittest.TestCase):
    """담당자는 자료를 압축으로 올린 뒤 제품 폴더에서 풀어 두는 일이 흔하다."""

    def setUp(self):
        self.root = tempfile.mkdtemp(prefix="pqr-dup-")
        inner = os.path.join(self.root, "2026 필요 자료")
        os.makedirs(inner)
        for name in ("8.2.1 RBG201 - 주원료 ERP.xls", "7. 수율현황표.xlsx"):
            with open(os.path.join(inner, name), "wb") as handle:
                handle.write(b"same-bytes")
        with zipfile.ZipFile(os.path.join(self.root, "2026 필요 자료.zip"), "w") as archive:
            for name in ("8.2.1 RBG201 - 주원료 ERP.xls", "7. 수율현황표.xlsx"):
                archive.writestr("2026 필요 자료/" + name, b"same-bytes")

    def tearDown(self):
        shutil.rmtree(self.root, ignore_errors=True)

    def test_폴더와_압축에_같은_파일이_있으면_한_번만(self):
        got = collect.discover(self.root)
        self.assertEqual(len(got.get("8.2.1", [])), 1)
        self.assertEqual(len(got.get("7", [])), 1)

    def test_남기는_것은_제품_폴더_안의_파일이다(self):
        # 담당자가 열어 고치는 것은 폴더 쪽이고, 압축 안의 것은 올릴 때 굳은 사본이다.
        got = collect.discover(self.root)
        self.assertTrue(os.path.abspath(got["8.2.1"][0]).startswith(os.path.abspath(self.root)))

    def test_크기가_다르면_다른_파일이다(self):
        with open(os.path.join(self.root, "2026 필요 자료", "8.2.1 다른 원료 ERP.xls"), "wb") as handle:
            handle.write(b"different-length-bytes")
        self.assertEqual(len(collect.discover(self.root).get("8.2.1", [])), 2)


def _table(rows, cols):
    d = docx.Document()
    t = d.add_table(rows=rows, cols=cols)
    return d, t


class 창에_자동으로_맞춤(unittest.TestCase):
    """전년도 결재본에서 물려받은 표는 본문 폭보다 넓어 오른쪽 여백을 넘는 것이 있었다."""

    def test_본문_폭에_딱_맞춘다(self):
        d, t = _table(2, 3)
        grid = t._tbl.find(qn("w:tblGrid"))
        for g, w in zip(grid.findall(qn("w:gridCol")), (2000, 4000, 4105)):
            g.set(qn("w:w"), str(w))
        self.assertTrue(E.fit_to_window(t, 9978))
        got = [int(g.get(qn("w:w"))) for g in grid.findall(qn("w:gridCol"))]
        self.assertEqual(sum(got), 9978)                      # 넘치지도 모자라지도 않는다
        self.assertLess(abs(got[0] / 9978.0 - 2000 / 10105.0), 0.01)   # 비율은 그대로

    def test_표_너비를_백_퍼센트로_둔다(self):
        d, t = _table(2, 2)
        E.fit_to_window(t, 9978)
        pr = t._tbl.find(qn("w:tblPr"))
        self.assertEqual(pr.find(qn("w:tblW")).get(qn("w:type")), "pct")
        self.assertEqual(pr.find(qn("w:tblW")).get(qn("w:w")), "5000")
        self.assertEqual(pr.find(qn("w:tblLayout")).get(qn("w:type")), "autofit")
        self.assertEqual(pr.find(qn("w:tblInd")).get(qn("w:w")), "0")

    def test_본문_폭을_모르면_손대지_않는다(self):
        d, t = _table(2, 2)
        self.assertFalse(E.fit_to_window(t, 0))

    def test_쪽_설정에서_본문_폭을_읽는다(self):
        d = docx.Document()
        sect = d.sections[0]._sectPr
        pg = sect.find(qn("w:pgSz")); pg.set(qn("w:w"), "11906")
        mar = sect.find(qn("w:pgMar")); mar.set(qn("w:left"), "1077"); mar.set(qn("w:right"), "851")
        self.assertEqual(E.text_width(d), 9978)


class 내용이_없으면(unittest.TestCase):
    """담당자 지시: 내용이 없으면 1행만 남겨두고 사선처리."""

    def test_빈_줄을_하나만_남긴다(self):
        d, t = _table(4, 3)
        t.rows[0].cells[0].text = "연번"
        self.assertEqual(E.collapse_empty_block(t, 1, 3), 2)
        self.assertEqual(len(t.rows), 2)

    def test_한_줄뿐이면_지우지_않는다(self):
        d, t = _table(2, 3)
        self.assertEqual(E.collapse_empty_block(t, 1, 1), 0)
        self.assertEqual(len(t.rows), 2)


if __name__ == "__main__":
    unittest.main()


class 목차_쪽수(unittest.TestCase):
    """Word 는 필드에 갱신 표시가 없으면 캐시된 옛 값을 그대로 보여 준다.

    담당자 화면에서 목차가 전부 '1' 로 나왔다 — 서식에 적혀 있던 남의 쪽수였다.
    LibreOffice 는 열 때마다 다시 계산해 PDF 로는 맞아 보여 놓치기 쉬웠다.
    """

    def _doc(self):
        d = docx.Document()
        toc = d.add_table(rows=2, cols=2)
        toc.rows[0].cells[0].text = "1. 목 적"
        toc.rows[0].cells[1].text = "24"          # 서식에 적혀 있던 옛 쪽수
        toc.rows[1].cells[0].text = "2. 배 경"
        toc.rows[1].cells[1].text = "24"
        d.add_paragraph("1. 목 적")
        d.add_paragraph("2. 배 경")
        return d, toc

    def test_필드에_갱신_표시를_붙인다(self):
        from pqr.engine.toc import link_toc
        d, toc = self._doc()
        self.assertEqual(link_toc(d, toc), 2)
        xml = d.element.body.xml
        self.assertEqual(xml.count('w:dirty="true"'), 2)
        self.assertIn("PAGEREF _pqr_toc_1", xml)
        self.assertIn("PAGEREF _pqr_toc_2", xml)

    def test_서식에_적혀_있던_옛_쪽수를_남기지_않는다(self):
        from pqr.engine.toc import link_toc
        d, toc = self._doc()
        link_toc(d, toc)
        self.assertNotIn("24", toc.rows[0].cells[1].text)

    def test_책갈피와_참조가_짝을_이룬다(self):
        import re
        from pqr.engine.toc import link_toc
        d, toc = self._doc()
        link_toc(d, toc)
        xml = d.element.body.xml
        names = set(re.findall(r'w:bookmarkStart[^>]*w:name="(_pqr_toc_[^"]+)"', xml))
        refs = set(re.findall(r"PAGEREF (_pqr_toc_\S+)", xml))
        self.assertEqual(names, refs)


class 비고_빈_칸(unittest.TestCase):
    """담당자 지시(2026-09): 3항 비고는 빈 줄끼리 합치고 사선 하나만 — 줄마다 N/A 를 적지 않는다."""

    def _doc(self):
        d = docx.Document()
        t = d.add_table(rows=5, cols=4)
        for k, v in enumerate(("No.", "점검 항목", "내  용", "비 고")):
            t.rows[0].cells[k].text = v
        t.rows[3].cells[3].text = "<주성분> 겐타마이신황산염"     # 가운데 한 줄만 내용이 있다
        return d, t

    def test_빈_줄끼리_합치고_사선_하나(self):
        d, t = self._doc()
        blocks, rows = E.merge_empty_runs(t, "비고")
        self.assertEqual((blocks, rows), (2, 3))               # 1~2행 묶음, 4행 묶음
        marks = []
        for tr in t._tbl.findall(qn("w:tr"))[1:]:
            tc = tr.findall(qn("w:tc"))[3]
            pr = tc.find(qn("w:tcPr"))
            vm = pr.find(qn("w:vMerge"))
            marks.append("-" if vm is None else (vm.get(qn("w:val")) or "continue"))
        self.assertEqual(marks, ["restart", "continue", "-", "restart"])

    def test_머리글_사이가_벌어져_있어도_찾는다(self):
        # 머리글이 '비 고' 처럼 사이가 벌어져 있는 표가 많다.
        d, t = self._doc()
        self.assertEqual(E.merge_empty_runs(t, "비고")[0], 2)

    def test_없는_열이면_손대지_않는다(self):
        d, t = self._doc()
        self.assertEqual(E.merge_empty_runs(t, "있을리없는열"), (0, 0))

    def test_합친_묶음의_첫_칸에만_사선(self):
        d, t = self._doc()
        E.merge_empty_runs(t, "비고")
        first = t._tbl.findall(qn("w:tr"))[1].findall(qn("w:tc"))[3]
        second = t._tbl.findall(qn("w:tr"))[2].findall(qn("w:tc"))[3]
        self.assertIsNotNone(first.find(qn("w:tcPr")).find(qn("w:tcBorders")).find(qn("w:tr2bl")))
        borders = second.find(qn("w:tcPr")).find(qn("w:tcBorders"))
        self.assertTrue(borders is None or borders.find(qn("w:tr2bl")) is None)
