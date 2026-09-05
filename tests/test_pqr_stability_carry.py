# -*- coding: utf-8 -*-
"""13항 안정성 — 올해 시험일지를 읽지 못하면 전년도 결재본에서 옮겨 온다.

담당자 2026-09: "안정성도 공란인데 전년도 PQR 결재본 참고해서 작성한 다음에
13항 최신 안정성 시험 파일로 업로드해서 작성하면 돼."
"""
import unittest

import docx

from pqr.engine import carry, docedit as E
from pqr.engine.recipe_ointment import _carry_131, _carry_133


def _merge(row, a, b):
    """가로 병합 — python-docx 의 merge 는 gridSpan 을 만들고 남는 칸을 없앤다."""
    return row.cells[a].merge(row.cells[b])


def _문서(제목, 표만들기):
    """항 제목 문단 하나와 그 아래 표 하나로 된 문서."""
    doc = docx.Document()
    doc.add_paragraph(제목)
    표만들기(doc)
    return doc


def _13_1(doc, 줄, 마지막열):
    t = doc.add_table(rows=1 + len(줄) + 1, cols=8)
    머리 = ["연번", "해당 연도", "시험 기간", "제조 번호", "포장 형태", "보관 조건", "완료 일자", 마지막열]
    for j, v in enumerate(머리):
        E.set_cell(t.rows[0].cells[j], v)
    for i, row in enumerate(줄, 1):
        for j, v in enumerate(row):
            E.set_cell(t.rows[i].cells[j], v)
    E.set_cell(t.rows[-1].cells[0], "특이사항 (Comment)\nN/A")
    return t


def _13_3(doc, 성분, 장기, 라벨):
    """시험항목 칸이 두 열을 덮는 경향 분석 표."""
    t = doc.add_table(rows=2 + len(장기) + len(라벨) + 1, cols=2 + len(성분))
    E.set_cell(_merge(t.rows[0], 0, 1), "시험항목")
    E.set_cell(_merge(t.rows[0], 1, len(성분)), "함량(%)")
    for k, name in enumerate(성분):
        E.set_cell(t.rows[1].cells[2 + k], name)
    for i, (해, 값) in enumerate(장기):
        row = t.rows[2 + i]
        E.set_cell(row.cells[0], "장기" if i == 0 else "")
        E.set_cell(row.cells[1], 해)
        for k, v in enumerate(값):
            E.set_cell(row.cells[2 + k], v)
    for i, (이름, 값) in enumerate(라벨):
        row = t.rows[2 + len(장기) + i]
        E.set_cell(_merge(row, 0, 1), 이름)
        for k, v in enumerate(값):
            E.set_cell(row.cells[2 + k], v)
    E.set_cell(t.rows[-1].cells[0], "특이사항 (Comment)\nN/A")
    return t


옛_13_1 = [["1", "2023(1)", "6M", "OGW701", "4.0g/Tube", "25±2℃", "2024.02.06", "진행중"],
            ["2", "2024", "Initial", "OGX901", "4.0g/Tube", "25±2℃", "2024.11.14", "진행중"]]
성분 = ["플루오로메톨론", "겐타마이신황산염"]
옛_장기 = [("2023(1)", ["98.2 ~ 99.5", "97.2 ~ 101.3"]), ("2024", ["99.6", "99.0"])]
옛_라벨 = [("관리규격", ["90.0 ~ 110.0", "90.0 ~ 120.0"]), ("최소", ["98.2", "97.2"]),
            ("최대", ["99.6", "101.3"]), ("경향분석 결과", ["적합", "적합"])]


class 전년도_13_1_옮기기(unittest.TestCase):
    def setUp(self):
        self.old = _문서("13.1 장기 안정성 시험",
                          lambda d: _13_1(d, 옛_13_1, "비고"))
        self.new = _문서("13.1 장기 안정성 시험",
                          lambda d: _13_1(d, [["", "", "", "", "", "", "", ""]] * 4, "실시 사유"))

    def test_열_이름이_같은_칸만_옮긴다(self):
        읽음 = carry.stability_tables(self.old)
        self.assertEqual(_carry_131(self.new, 읽음["13.1"], {}, "PQR25.doc"), 1)
        t = self.new.tables[0]
        글 = [[E.cell_text(c) for c in E.raw_cells(r)] for r in t.rows]
        self.assertEqual(글[1][:8],
                         ["1", "2023(1)", "6M", "OGW701", "4.0g/Tube", "25±2℃", "2024.02.06", ""])
        self.assertEqual(글[2][3], "OGX901")

    def test_실시_사유는_밸리데이션_사유에서_채운다(self):
        읽음 = carry.stability_tables(self.old)
        _carry_131(self.new, 읽음["13.1"], {"OGW701": "Lot size 축소"}, "PQR25.doc")
        t = self.new.tables[0]
        self.assertEqual(E.cell_text(E.raw_cells(t.rows[1])[7]), "Lot size 축소")

    def test_어디서_옮겼는지_특이사항에_남긴다(self):
        읽음 = carry.stability_tables(self.old)
        _carry_131(self.new, 읽음["13.1"], {}, "PQR25.doc")
        꼬리 = E.cell_text(E.raw_cells(self.new.tables[0].rows[-1])[0])
        self.assertIn("전년도 결재본", 꼬리)
        self.assertIn("PQR25.doc", 꼬리)


class 전년도_13_3_옮기기(unittest.TestCase):
    def setUp(self):
        self.old = _문서("13.2 안정성 시험 경향 분석 결과",
                          lambda d: _13_3(d, 성분, 옛_장기, 옛_라벨))
        빈장기 = [("", ["", ""]), ("", ["", ""])]
        빈라벨 = [(이름, ["", ""]) for 이름, _ in 옛_라벨]
        self.new = _문서("13.3 안정성 시험 경향 분석 결과",
                          lambda d: _13_3(d, 성분, 빈장기, 빈라벨))

    def test_성분_이름으로_열을_맞춘다(self):
        읽음 = carry.stability_tables(self.old)
        self.assertEqual(_carry_133(self.new, 읽음["13.3"], "PQR25.doc"), 1)
        글 = [[E.cell_text(c) for c in E.raw_cells(r)] for r in self.new.tables[0].rows]
        self.assertEqual(글[2][1:4], ["2023(1)", "98.2 ~ 99.5", "97.2 ~ 101.3"])
        self.assertEqual(글[3][1:4], ["2024", "99.6", "99.0"])

    def test_라벨_줄은_병합_칸_옆부터_채운다(self):
        """'관리 규격' 칸이 두 열을 덮는다 — 자리로 세면 값이 한 칸 밀린다."""
        읽음 = carry.stability_tables(self.old)
        _carry_133(self.new, 읽음["13.3"], "PQR25.doc")
        줄 = {E.cell_text(E.raw_cells(r)[0]): [E.cell_text(c) for c in E.raw_cells(r)[1:]]
              for r in self.new.tables[0].rows}
        self.assertEqual(줄["관리규격"], ["90.0 ~ 110.0", "90.0 ~ 120.0"])
        self.assertEqual(줄["최소"], ["98.2", "97.2"])
        self.assertEqual(줄["경향분석 결과"], ["적합", "적합"])


if __name__ == "__main__":
    unittest.main()
