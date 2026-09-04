# -*- coding: utf-8 -*-
"""전년도 결재본(.doc)을 바꾸지 못할 때 — 요약본으로 떨어지지 않고 EDMS 빈 서식으로 만든다.

담당자 PC 에서 Word COM 이 막혀 옛 워드(.doc)를 .docx 로 바꾸지 못하는 일이 있다(2026-09,
디겐타안연고). 예전에는 여기서 멈춰 결재본 양식이 아닌 요약본이 나왔다. 이제는 서식은 지키고
전년도 결재본이 있어야만 채울 수 있는 항을 문의로 남긴다.
"""
import datetime
import os
import shutil
import tempfile
import unittest

import docx

from pqr.engine import convert, writer


def _doc(rows):
    """머리행 하나에 rows(각각 칸 글 목록)를 붙인 표 하나짜리 문서."""
    d = docx.Document()
    d.add_paragraph("9.1 시험결과표")
    t = d.add_table(rows=1 + len(rows), cols=3)
    for k, v in enumerate(("연번", "시험항목", "결과")):
        t.rows[0].cells[k].text = v
    for i, row in enumerate(rows):
        for k, v in enumerate(row):
            t.rows[1 + i].cells[k].text = v
    return d


class 빈표찾기(unittest.TestCase):
    def test_값이_없으면_항_번호를_돌려준다(self):
        self.assertEqual(writer.blank_sections(_doc([["1", "", ""], ["2", "", ""]])), ["9.1"])

    def test_값이_있으면_비지_않았다(self):
        self.assertEqual(writer.blank_sections(_doc([["1", "성상", "적합"]])), [])

    def test_연번만_있는_것은_값이_아니다(self):
        self.assertEqual(writer.blank_sections(_doc([["1", "", ""], ["2", "", ""]])), ["9.1"])

    def test_서식에_원래_있는_표시는_값이_아니다(self):
        self.assertEqual(writer.blank_sections(_doc([["1", "", "☐ Yes ■ No"]])), ["9.1"])

    def test_해당없음은_값이다(self):
        # 담당자가 '해당 없음' 이라고 쓴 표는 채워진 표다 — 비었다고 알리면 안 된다.
        self.assertEqual(writer.blank_sections(_doc([["1", "해당 없음", ""]])), [])

    def test_제목이_없는_표는_따지지_않는다(self):
        d = docx.Document()
        t = d.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "연번"
        self.assertEqual(writer.blank_sections(d), [])


class 변환실패하면(unittest.TestCase):
    def setUp(self):
        self.root = tempfile.mkdtemp(prefix="pqr-fallback-")
        self.folder = os.path.join(self.root, "QC1-9999 시험제품")
        os.makedirs(self.folder)
        with open(os.path.join(self.folder, "16. PQR25 시험제품.doc"), "wb") as h:
            h.write(b"\xd0\xcf\x11\xe0" + b"0" * 64)         # 옛 워드처럼 보이는 파일
        self.real = convert.to_docx

        def fake(src, dst):
            if src.lower().endswith(".doc"):
                raise convert.ConvertError("시험용 실패")
            return self.real(src, dst)
        convert.to_docx = fake

    def tearDown(self):
        convert.to_docx = self.real
        shutil.rmtree(self.root, ignore_errors=True)

    def test_요약본으로_떨어지지_않고_EDMS_서식으로_만든다(self):
        out = os.path.join(self.root, "out.docx")
        got = writer.write_report(self.folder,
                                  {"code": "QC1-9999", "name": "시험제품", "group": "안연고"},
                                  {"from": "2025-01-01", "to": "2025-12-31"}, out,
                                  today=datetime.date(2026, 9, 4))
        self.assertTrue(os.path.isfile(out))
        d = docx.Document(out)
        texts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        self.assertEqual(texts[0], "목차 (Table of Contents)")
        self.assertGreater(len(d.tables), 25)                 # 결재본 양식 그대로
        kinds = [i[0] for i in got["issues"]]
        self.assertIn("16", kinds)                            # 전년도 결재본을 못 읽었다고 알린다
        self.assertIn("서식", kinds)                          # 채우지 못한 항을 알린다
        why = next(i[2] for i in got["issues"] if i[0] == "16")
        self.assertIn("Word 문서(*.docx)", why)               # 담당자가 할 일을 알려 준다

    def test_서식도_없으면_그대로_알린다(self):
        # 서식을 찾을 길이 아예 없으면 예전처럼 오류로 알린다 — 조용히 엉뚱한 문서를 내지 않는다.
        from pqr.engine import edms
        real_find = edms.find_form
        edms.find_form = lambda folder, depth=2: None
        try:
            with self.assertRaises(convert.ConvertError):
                writer.write_report(self.folder,
                                    {"code": "QC1-9999", "name": "시험제품", "group": "안연고"},
                                    {"from": "2025-01-01", "to": "2025-12-31"},
                                    os.path.join(self.root, "out2.docx"),
                                    today=datetime.date(2026, 9, 4))
        finally:
            edms.find_form = real_find


if __name__ == "__main__":
    unittest.main()
