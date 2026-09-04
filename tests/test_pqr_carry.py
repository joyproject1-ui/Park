# -*- coding: utf-8 -*-
"""전년도 결재본에서 이어받기 — 해마다 바뀌지 않는 값만, 빈 칸에만.

담당자 지시(2026-09): "전년도 PQR 결재본은 장비 및 원료 등 정보를 참고하는거고, 2026년 PQR
필요 자료를 참고해서 새롭게 작성해줘야 하는거야."
"""
import unittest

import docx

from pqr.engine import carry


def _doc(rows, heading="8.1.1 주원료"):
    document = docx.Document()
    document.add_paragraph(heading)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for ri, cells in enumerate(rows):
        for ci, text in enumerate(cells):
            table.cell(ri, ci).text = text
    return document


HEAD = ["연번", "관리번호", "원/자재명", "규격", "제조원", "완료일"]


class 빈_칸에만_이어받는다(unittest.TestCase):
    def test_같은_코드의_규격을_가져온다(self):
        old = _doc([HEAD, ["1", "RBG201", "겐타마이신황산염", "KP", "예전 제조원", "2023.12.08"]])
        new = _doc([HEAD, ["1", "RBG201", "겐타마이신황산염", "", "올해 제조원", ""]])
        carry.carry(new, old)
        got = [c.text for c in new.tables[0].rows[1].cells]
        self.assertEqual(got[3], "KP")                 # 규격은 이어받고
        self.assertEqual(got[4], "올해 제조원")          # 올해 자료로 채운 칸은 그대로 두고
        self.assertEqual(got[5], "")                   # 완료일은 해마다 바뀌므로 이어받지 않는다

    def test_코드가_다르면_가져오지_않는다(self):
        old = _doc([HEAD, ["1", "RSF101", "플루오로메톨론", "USP", "", ""]])
        new = _doc([HEAD, ["1", "RBG201", "겐타마이신황산염", "", "", ""]])
        carry.carry(new, old)
        self.assertEqual([c.text for c in new.tables[0].rows[1].cells][3], "")


class 줄_열쇠가_해마다_다른_표(unittest.TestCase):
    """6항은 Lot No. 가 해마다 달라 짝이 없다 — 전년도 값이 한 가지면 그 값을 쓴다."""

    HEAD6 = ["연번", "Lot No.", "제조일자", "제조단위", "포장단위"]

    def test_제조단위는_전년도_값이_하나면_가져온다(self):
        old = _doc([self.HEAD6, ["1", "OGX901", "2024.09.24", "164,000g", "4g x 1Tube/Case"]],
                   "6. 제조내역 확인")
        new = _doc([self.HEAD6, ["1", "OGY301", "2025.03.11", "", ""]], "6. 제조내역 확인")
        carry.carry(new, old)
        got = [c.text for c in new.tables[0].rows[1].cells]
        self.assertEqual(got[3], "164,000g")
        self.assertEqual(got[4], "4g x 1Tube/Case")
        self.assertEqual(got[1], "OGY301")            # Lot 은 올해 것 그대로

    def test_전년도_값이_여럿이면_고르지_않는다(self):
        old = _doc([self.HEAD6,
                    ["1", "OGX901", "2024.09.24", "164,000g", "4g x 1Tube/Case"],
                    ["2", "OGX902", "2024.10.24", "82,000g", "4g x 1Tube/Case"]], "6. 제조내역 확인")
        new = _doc([self.HEAD6, ["1", "OGY301", "2025.03.11", "", ""]], "6. 제조내역 확인")
        carry.carry(new, old)
        self.assertEqual([c.text for c in new.tables[0].rows[1].cells][3], "")


if __name__ == "__main__":
    unittest.main()
