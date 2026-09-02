# -*- coding: utf-8 -*-
"""결재본 안에서 표를 '항 제목' 으로 찾는다 (표 번호로 찾지 않는다).

제품마다 결재본의 표 개수와 순서가 달라 T[38] 같은 번호는 다른 제품에서 엉뚱한 표를
가리킨다. 표 바로 앞에 나오는 제목('11.1.1. 내수용')과 표의 머리행 글자로 찾아야 한다.
"""
import re
from docx.oxml.ns import qn


def _text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


HEADING = re.compile(r"^\d{1,2}(\.\d+)*\.?\s*\S")


def outline(document):
    """본문 순서대로 [('h', 제목 글, 문단 el) | ('t', 표 번호, 표 el)]"""
    body = document.element.body
    tables = [t._tbl for t in document.tables]
    out = []
    for el in body:
        if el.tag == qn("w:tbl"):
            out.append(("t", tables.index(el), el))
        elif el.tag == qn("w:p"):
            text = _text(el)
            if text and HEADING.match(text) and len(text) < 80:
                out.append(("h", text, el))
    return out


def headings_of_tables(document):
    """{표 번호: [바로 앞 제목들 (가까운 순)]} — 표 사이에 제목이 여럿이면 모두."""
    result, stack = {}, []
    for kind, value, _ in outline(document):
        if kind == "h":
            stack.append(value)
        else:
            result[value] = list(reversed(stack))
            stack = []
    return result


def find_tables(document, heading_prefix, header_words=(), after=None):
    """heading_prefix 로 시작하는 제목 뒤에 오는 표들 (header_words 가 머리행에 모두 있으면).

    같은 제목 아래 표가 여럿(9.2.1.1 성상표·이물표 …)이면 순서대로 돌려준다.
    """
    hits, current = [], False
    key = re.sub(r"\s+", "", heading_prefix)
    for kind, value, el in outline(document):
        if kind == "h":
            current = re.sub(r"\s+", "", value).startswith(key)
            continue
        if not current:
            continue
        table = document.tables[value]
        if after is not None and value <= after:
            continue
        if header_words:
            head = _text(table._tbl.findall(qn("w:tr"))[0])
            if not all(w in re.sub(r"\s+", "", head) for w in header_words):
                continue
        hits.append(value)
    return hits


def find_table(document, heading_prefix, header_words=(), nth=0):
    hits = find_tables(document, heading_prefix, header_words)
    if len(hits) <= nth:
        raise KeyError("표를 찾지 못했습니다: 제목 '%s' 머리 %s (%d번째)" % (heading_prefix, header_words, nth))
    return document.tables[hits[nth]], hits[nth]


def find_para(document, needle):
    for p in document.paragraphs:
        if needle in p.text:
            return p
    return None
