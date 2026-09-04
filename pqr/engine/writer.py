# -*- coding: utf-8 -*-
"""'보고서 작성' 한 번에 결재본 양식의 제출용 보고서를 만든다.

절차: 전년도 결재본 → .docx 로 → 자료 판독(collect) → 항별 채움(recipe) → 조판(layout)
→ 저장 → 변환 흔적·OOXML 순서 정리(polish) → 순서 검사 0 건 확인.
헷갈린 값은 issues(문의 목록)로 남기고 보고서에는 '확인 필요' 표시를 둔다.
"""
import os
import re
import shutil
import tempfile
import zipfile

from . import collect as collect_module
from . import convert, layout, polish
from .polish import check_order


class EngineError(Exception):
    pass


def _year_of(text):
    m = re.search(r"PQR(\d{2})-", text or "")
    return (2000 + int(m.group(1))) if m else None


def find_previous(folder, workdir=None):
    """평가항목 16 의 전년도 결재본. .doc/.docx 를 먼저, 없으면 .zip 안의 것을 workdir 에 꺼내 돌려준다.

    담당자는 전년도 PQR 을 워드와 Cpk 엑셀을 함께 압축해 올리는 일이 잦다(퀴노비드·디겐타 모두
    '16 전년도 PQR word & excel (PQR25).zip'). 압축을 안 보면 엔진이 "결재본 없음" 으로 멈추고
    예전 방식은 .docx 만 찾아 요약본이 나온다 — 디겐타안연고에서 실제로 생겼다.
    """
    root = os.path.abspath(folder)
    data_files = collect_module.discover(folder, workdir)
    files = [p for p in data_files.get("16", []) if not os.path.basename(p).startswith("~$")]

    def rank(p):
        inside = not os.path.abspath(p).startswith(root)      # 판독기가 압축에서 꺼낸 것
        return (inside, 0 if p.lower().endswith(".docx") else 1, p)
    docs = sorted((p for p in files if p.lower().endswith((".docx", ".doc"))), key=rank)
    if docs:
        return docs[0]
    for p in files:
        if p.lower().endswith(".zip"):
            got = _doc_from_zip(p, workdir or tempfile.mkdtemp(prefix="pqr-prev-"))
            if got:
                return got
    return None


def _doc_from_zip(zip_path, workdir):
    """압축 안의 워드(.docx 우선, 없으면 .doc — 여럿이면 가장 큰 것)를 꺼낸다."""
    with zipfile.ZipFile(zip_path) as z:
        cands = []
        for info in z.infolist():
            name = info.filename
            try:
                name = name.encode("cp437").decode("cp949")
            except Exception:
                pass
            base = os.path.basename(name)
            if base.startswith(("~$", ".")) or "__MACOSX" in name:
                continue
            if base.lower().endswith((".docx", ".doc")):
                cands.append((0 if base.lower().endswith(".docx") else 1, -info.file_size, info, base))
        if not cands:
            return None
        cands.sort()
        _, _, info, base = cands[0]
        target = os.path.join(workdir, base)
        with open(target, "wb") as h:
            h.write(z.read(info.filename))
        return target


def _plain(text):
    return re.sub(r"[\s()（）\[\]·ㆍ・,.-]", "", text or "")


def other_product(base_docx, product, shown=None):
    """전년도 결재본이 이 제품 것이 아니면 그 이유 문장을, 맞으면 None 을 돌려준다.

    제품명(계획서 이름)의 앞 네 글자나 제품코드가 머리글·본문 앞부분에 있어야 한다.
    '퀴노비드안연고' 결재본으로 '디겐타안연고' 를 만드는 일을 막는다(담당자 지적 2026-09-04).
    """
    name = _plain(product.get("name") or "")
    code = _plain(product.get("code") or "").upper()      # 하이픈을 뺀 꼴로 견준다 (QC1-7014 → QC17014)
    if not name and not code:
        return None
    try:
        import docx
        d = docx.Document(base_docx)
    except Exception:
        return None
    parts = []
    for section in d.sections:
        for part in (section.header, section.first_page_header):
            try:
                parts.append(" ".join(p.text for p in part.paragraphs))
                parts.append(" ".join(E_cell(c) for t in part.tables for r in t.rows for c in r.cells))
            except Exception:
                pass
    parts.append(" ".join(p.text for p in d.paragraphs[:120]))
    parts.append(" ".join(c.text for t in d.tables[:4] for r in t.rows for c in r.cells))
    hay = _plain(" ".join(parts))
    key = name[:4]
    if (key and key in hay) or (code and code in hay.upper()):
        return None
    return ("전년도 결재본(%s)에 이 제품(%s %s)의 이름이 없습니다 — 다른 제품의 결재본으로 보입니다. "
            "제품 폴더의 '16. 전년도 PQR' 파일을 확인하세요."
            % (shown or os.path.basename(base_docx), (product.get("code") or "").strip(), product.get("name") or ""))


def E_cell(cell):
    return cell.text


def attachment_source(folder, workdir=None):
    """첨부 Cpk 엑셀을 물려받을 곳 — 항목 16 의 .zip, 아니면 결재본(과 같은 폴더의 .xls)."""
    data_files = collect_module.discover(folder, workdir)
    for p in data_files.get("16", []):
        if p.lower().endswith(".zip"):
            return p
    return find_previous(folder, workdir)


BLANK_MARKS = re.compile(r"[☐■□☑✔]|Yes|No")     # 서식에 원래 있는 표시 — 값이 아니다


def _row_text(row_el):
    from docx.oxml.ns import qn
    cells = row_el.findall(qn("w:tc"))
    out = []
    for tc in cells[1:]:                      # 연번 칸(1,2,3…)은 서식에 이미 있다
        out.append("".join(t.text or "" for t in tc.iter(qn("w:t"))))
    text = " ".join(out)
    return BLANK_MARKS.sub("", text).strip()


def blank_sections(document):
    """값이 하나도 안 들어간 표의 항 번호 목록. EDMS 빈 서식만으로 만들었을 때 무엇이 비었는지 알린다.

    전년도 결재본이 있으면 표 구조(제품마다 다른 시험항목·설비 목록)가 거기서 오지만, 빈 서식만
    있으면 채울 자리가 없어 그대로 빈 표가 남는다. 담당자가 그 항을 직접 써야 하므로 짚어 준다.
    """
    from docx.oxml.ns import qn
    from .locate import headings_of_tables
    heads = headings_of_tables(document)
    blank = []
    for i, table in enumerate(document.tables):
        titles = heads.get(i) or []
        head = next((t for t in titles if re.match(r"^\d", t)), None)
        if not head:
            continue
        rows = table._tbl.findall(qn("w:tr"))
        if len(rows) < 2:
            continue
        if any(_row_text(r) for r in rows[1:]):
            continue
        number = re.match(r"^\d{1,2}(\.\d+)*", head).group(0).rstrip(".")
        if number not in blank:
            blank.append(number)
    return blank


def write_report(folder, product, period, out_path, today=None, recipe=None, log=None, vision=None,
                 ignore_previous=False):
    """folder: 제품 폴더 · product: {"code","name","group"} · period: {"from","to"} · out_path: 저장할 .docx

    ignore_previous: 전년도 결재본을 쓰지 않고 EDMS 빈 서식만으로 만든다. 전년도 결재본을 바탕
    으로 한 첫 시도가 넘어졌을 때 부르는 쪽이 한 번 더 해 보는 길이다 — 결재본 양식은 지켜야
    하므로, 자료 상태 요약본으로 물러나기 전에 이것을 먼저 해 본다.

    돌려주는 값: {"path", "issues": [(항, 파일, 설명)], "log": [...], "blank_sections": [...]}
    """
    lines = []
    blank = []

    def log_(msg):
        lines.append(msg)
        if log:
            log(msg)

    # 본문은 전년도 결재본을 바탕으로 채운다 — 제품마다 다른 항·표 구조와 문안이 거기 있다.
    # 서식은 언제나 EDMS 결재본 서식(E-HLF-32): 채운 본문을 서식 껍데기에 옮겨 담는다(rehouse).
    # 전년도 결재본이 없는 첫해 제품은 서식을 바탕으로 직접 채운다(표 구조는 서식 그대로).
    from . import edms, rehouse
    work = tempfile.mkdtemp(prefix="pqr-report-")
    previous = None if ignore_previous else find_previous(folder, work)
    form = edms.find_form(folder)
    data_issue = None
    # 제품용 최신 양식(EDMS 서식이면서 이 제품의 항·표·허용기준이 든 빈 양식)이 폴더에 있으면
    # 그것이 바탕이다 — 담당자 지시(2026-09): "올해부터 최신양식으로 변경됐어."
    # 전년도 결재본은 옛 양식이라 거기 없는 항(디겐타안연고 9.2 세부 시험결과)이 통째로 빠진다.
    product_form = None
    if form and not edms.is_shipped(form):
        if edms.is_named_form(form):
            # 담당자가 평가항목 (v) 'PQR 작성 공양식' 으로 올린 양식 — 제품명이 안 적힌 빈 양식일
            # 수 있으니 이름 검사 대신, 남의 제품코드가 적혀 있는지만 본다.
            wrong = edms.mentions_other_code(form, product.get("code"))
            if wrong:
                raise EngineError("공양식(%s)에 다른 제품코드 %s 가 적혀 있습니다 — 이 제품(%s)의 "
                                  "양식이 맞는지 확인하세요." % (os.path.basename(form), wrong,
                                                            product.get("code") or ""))
            product_form = form
        elif other_product(form, product) is None:
            product_form = form
    if product_form:
        source, why = product_form, "EDMS 최신 양식(이 제품용)"
    elif previous:
        source, why = previous, "전년도 결재본"
    elif form:
        source, why = form, "EDMS 서식(전년도 결재본 없음 — 첫해 제품)"
    else:
        raise EngineError(edms.choose_base(None, None)[1])
    base = os.path.join(work, "base.docx")
    try:
        how = convert.to_docx(source, base)
    except convert.ConvertError as error:
        # 옛 워드(.doc)를 바꿀 길이 없는 PC 가 있다(Word COM 이 막힌 회사 PC). 예전에는 여기서
        # 멈춰 결재본 양식 대신 요약본이 나왔다. 이제는 EDMS 빈 서식으로 바탕을 바꿔 결재본
        # 양식은 지키고, 전년도 결재본에서만 얻을 수 있는 것(전년도 시정사항 등)을 문의로 남긴다.
        if source is not previous or not form:
            raise
        log_("전년도 결재본을 읽지 못해 EDMS 빈 서식으로 만듭니다 — %s" % error)
        shown = os.path.basename(previous)
        source, why = form, "EDMS 서식(전년도 결재본을 .docx 로 바꾸지 못함)"
        previous = None
        how = convert.to_docx(source, base)
        data_issue = ("16", shown,
                      "전년도 결재본(.doc)을 읽지 못해 EDMS 빈 서식으로 만들었습니다 — "
                      "16항 전년도 시정사항과 전년도 문안을 직접 확인하세요. "
                      "그 파일을 Word 에서 '다른 이름으로 저장 → Word 문서(*.docx)' 로 저장해 "
                      "제품 폴더에 두면 다음부터는 전년도 결재본을 그대로 물려받습니다.")
    log_("바탕 문서: %s — %s (%s)" % (os.path.basename(source), why, how))
    if source is previous:
        # 다른 제품의 결재본으로 쓰면 제품명·규격·표 구성이 통째로 남의 것이 된다 — 여기서 막는다.
        wrong = other_product(base, product, os.path.basename(source))
        if wrong:
            raise EngineError(wrong)
    if product_form:
        log_("이 제품의 EDMS 최신 양식을 바탕으로 씁니다 — 전년도 결재본은 참고만 합니다")
    elif form and source is previous:
        log_("EDMS 서식: %s — 채운 뒤 이 서식에 옮겨 담음" % (
            "프로그램에 든 E-HLF-32 빈 서식" if edms.is_shipped(form) else os.path.basename(form)))
    elif not form:
        log_("EDMS 서식 없음 — 제품 폴더나 공통 폴더에 E-HLF-32 서식(.docx)을 두면 그 서식으로 만듭니다")
        data_issue = ("서식", "", "EDMS 결재본 서식(E-HLF-32)이 없어 전년도 양식 그대로 만들었음 — 제품 폴더나 '공통' 폴더에 서식을 두세요")

    data = collect_module.collect(folder, product_name=product.get("name"), log=log_)
    data.previous_report = previous
    data.period = period
    if vision is not None:
        try:
            vision(data, log_)
        except Exception as error:               # 비전 판독 실패는 보고서를 막지 않는다
            data.issues.append(("13", "", "손글씨 판독 실패: %s" % error))

    import docx
    document = docx.Document(base)
    if recipe is None:
        from . import recipe_ointment as recipe_module
        recipe = recipe_module.fill
    if product_form and previous:
        # 최신 양식을 바탕으로 쓸 때 전년도 결재본은 참고 자료다 — 해마다 바뀌지 않는 값
        # (원료 규격·제조단위·포장단위·제조원)만 빈 칸에 먼저 이어받고, 그 위에 올해 자료를 덮는다.
        from . import carry as carry_module
        try:
            prev_docx = os.path.join(work, "prev.docx")
            convert.to_docx(previous, prev_docx)
            old_document = docx.Document(prev_docx)
            carry_module.carry(document, old_document, log_)
            data.pv_reasons = carry_module.pv_reasons(old_document)
            if data.pv_reasons:
                log_("전년도 10.1 에서 밸리데이션 사유 %d Lot" % len(data.pv_reasons))
        except Exception as error:
            log_("전년도 결재본을 참고하지 못했습니다 — %s" % error)
            data.issues.append(("16", os.path.basename(previous),
                                "전년도 결재본을 읽지 못해 원료 규격·제조단위 같은 값을 이어받지 "
                                "못했습니다 — 빈 칸을 직접 채우세요"))
    ctx = recipe(document, data, product, period, today=today, log=log_)
    if form and source is previous:
        filled = os.path.join(work, "filled.docx")
        document.save(filled)
        housed = os.path.join(work, "housed.docx")
        got = rehouse.rehouse(filled, form, housed)
        log_("EDMS 서식에 옮겨 담음: %s" % ", ".join("%s %s" % kv for kv in got.items()))
        document = docx.Document(housed)
    elif data_issue:
        data.issues.append(data_issue)
    if previous is None:
        # 전년도 결재본 없이 빈 서식만으로 만들면 표 구조가 없어 못 채우는 항이 남는다 — 짚어 준다.
        empty = blank_sections(document)
        blank[:] = empty
        if empty:
            log_("빈 서식만으로 만들어 채우지 못한 항: %s" % ", ".join(empty))
            data.issues.append(("서식", "", "%s 항의 표를 채우지 못했습니다 — 이 항들은 제품마다 표 구성이 "
                                "달라 전년도 결재본에서 표를 물려받아야 합니다. 직접 작성하시거나, "
                                "전년도 결재본을 .docx 로 제품 폴더에 두고 다시 만드세요."
                                % ", ".join(empty)))
    layout.apply(document, log=log_, product_title=(ctx or {}).get("cover_title"))
    document.save(out_path)
    polish.polish(out_path)
    with zipfile.ZipFile(out_path) as z:
        bad = sum(check_order(z.read(n).decode("utf-8")) for n in z.namelist()
                  if n.startswith("word/") and n.endswith(".xml"))
    if bad:
        raise EngineError("OOXML 순서 검사에 실패했습니다 (%d 곳)" % bad)
    from . import verify                           # Word 가 '복구할까요' 를 묻는 문제 마지막 검사
    problems = verify.check_docx(out_path)
    if problems:
        raise EngineError("Word 가 열지 못할 문서입니다 — " + "; ".join(problems[:5]))
    log_("문서 검사 통과 (부품·관계·스타일·요소 차례)")
    log_("순서 검사 0 건 · 저장: %s" % os.path.basename(out_path))
    # 목차 쪽수·머리글 Page 는 Word 가 쪽을 나눠 봐야 안다. 담당자 PC 에는 Word 가 있으므로
    # 여기서 한 번 계산해 둔다 — 그러지 않으면 파일을 열었을 때 목차가 모두 '1' 로 보인다.
    if convert.refresh_fields(out_path):
        log_("목차 쪽수·머리글을 Word 로 다시 계산했습니다")
    else:
        log_("Word 가 없어 목차 쪽수는 파일을 열 때 계산됩니다 — 숫자가 모두 1 이면 "
             "Ctrl+A 를 누르고 F9 를 누르세요")
    # 첨부 엑셀 — Cpk 계산 파일 4종(결재본 것을 물려받아 값 갱신) + 안정성 경향 분석
    attachments = []
    try:
        from . import excel_attach
        day = today.strftime("%Y.%m.%d") if hasattr(today, "strftime") else re.sub(r"-", ".", str(today or ""))[:10]
        attachments += excel_attach.write_cpk_files(os.path.dirname(out_path), data,
                                                    attachment_source(folder, work) or previous, day)
        attachments += excel_attach.write_stability_workbook(
            os.path.dirname(out_path), data, product, day,
            input_dir=os.path.dirname(os.path.abspath(folder)), report_path=out_path,
            product_dir=folder)
        log_("첨부 엑셀: %s" % ", ".join(n for n, _ in attachments))
    except Exception as error:
        data.issues.append(("첨부", "", "첨부 엑셀 생성 실패: %s" % error))
    shutil.rmtree(work, ignore_errors=True)
    return {"path": out_path, "issues": data.issues + list((ctx or {}).get("issues", [])), "log": lines,
            "data": data, "attachments": attachments, "blank_sections": blank}
