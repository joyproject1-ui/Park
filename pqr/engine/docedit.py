# -*- coding: utf-8 -*-
"""전년도 PQR(.docx) 을 그대로 두고 값만 갈아 끼우기 위한 도우미."""
import copy
import math
import re

from docx.oxml.ns import qn
from lxml import etree
from .ooxml_order import place, get_or_add, resort_all

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ---------- 문단/런 ----------
def _first_run(p):
    for r in p.findall(qn("w:r")):
        if r.find(qn("w:t")) is not None:
            return r
    return None


def set_para_text(p, text):
    """문단의 글자만 바꾼다. 첫 런의 서식을 유지하고 나머지 런의 글자는 지운다."""
    p = getattr(p, "_p", p)
    runs = [r for r in p.findall(qn("w:r")) if r.find(qn("w:t")) is not None]
    if not runs:
        if text:
            r = p.makeelement(qn("w:r"), {})
            t = p.makeelement(qn("w:t"), {})
            t.text = text
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            p.append(r)
        return
    first = True
    for r in runs:
        for t in r.findall(qn("w:t")):
            if first:
                t.text = text
                t.set(qn("xml:space"), "preserve")
                first = False
            else:
                t.text = ""


def para_text(p):
    return "".join(t.text or "" for t in p.iter(qn("w:t")))


# ---------- 셀 ----------
def cell_paras(tc):
    return tc.findall(qn("w:p"))


def set_cell(cell, *lines):
    """셀 글자를 줄 단위로 바꾼다. 문단 수가 모자라면 첫 문단을 복제한다."""
    tc = cell._tc
    paras = cell_paras(tc)
    lines = list(lines) or [""]
    while len(paras) < len(lines):
        new = copy.deepcopy(paras[0])
        paras[-1].addnext(new)
        paras = cell_paras(tc)
    for i, p in enumerate(paras):
        set_para_text(p, lines[i] if i < len(lines) else "")


def cell_text(cell):
    return "\n".join(para_text(p) for p in cell_paras(cell._tc))


# ---------- 행 ----------
def clone_row(table, src_index, count):
    """src_index 행을 복제해 그 뒤에 count 개를 넣는다. 복제본 리스트를 돌려준다."""
    src = table.rows[src_index]._tr
    made = []
    anchor = src
    for _ in range(count):
        new = copy.deepcopy(src)
        anchor.addnext(new)
        anchor = new
        made.append(new)
    return made


def drop_row(table, index):
    tr = table.rows[index]._tr
    tr.getparent().remove(tr)


def fit_rows(table, first_data, last_data, want):
    """first_data..last_data(포함) 구간의 데이터 행 수를 want 개로 맞춘다."""
    have = last_data - first_data + 1
    if want > have:
        clone_row(table, last_data, want - have)
    elif want < have:
        for _ in range(have - want):
            drop_row(table, last_data)
            last_data -= 1
    return first_data, first_data + want - 1


# ---------- 병합 ----------
def _tcpr(tc):
    pr = tc.find(qn("w:tcPr"))
    if pr is None:
        pr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, pr)
    return pr


def set_vmerge(cell, value):
    """value: 'restart' 이면 시작, None 이면 이어짐, False 면 병합 해제."""
    tc = cell._tc
    pr = _tcpr(tc)
    old = pr.find(qn("w:vMerge"))
    if old is not None:
        pr.remove(old)
    if value is False:
        return
    vm = pr.makeelement(qn("w:vMerge"), {})
    if value == "restart":
        vm.set(qn("w:val"), "restart")
    pr.append(vm)


# ---------- 사선 ----------
def clear_diag(cell):
    pr = _tcpr(cell._tc)
    borders = pr.find(qn("w:tcBorders"))
    if borders is None:
        return
    for tag in ("w:tl2br", "w:tr2bl"):
        el = borders.find(qn(tag))
        if el is not None:
            borders.remove(el)


def add_diag(cell):
    """회사 관행: 왼쪽 아래 → 오른쪽 위 사선 하나."""
    pr = _tcpr(cell._tc)
    borders = get_or_add(pr, "tcBorders")
    for tag in ("w:tl2br", "w:tr2bl"):
        el = borders.find(qn(tag))
        if el is not None:
            borders.remove(el)
    el = get_or_add(borders, "tr2bl")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "4")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")


def has_diag(cell):
    pr = cell._tc.find(qn("w:tcPr"))
    borders = pr.find(qn("w:tcBorders")) if pr is not None else None
    return borders is not None and any(borders.find(qn(t)) is not None for t in ("w:tl2br", "w:tr2bl"))


NA_ONLY = re.compile(r"^\s*(N\s*/\s*A|해당\s*없음|없음)\s*$", re.I)


def drop_na_in_diag_cells(document):
    """사선을 그은 칸에 남아 있는 'N/A' 를 지운다. 지운 칸 수를 돌려준다.

    담당자 지시(2026-09): "비고에 사선이 그어졌으니 N/A 는 삭제해줘." 사선이 곧 '해당 없음'
    이라 글자를 겹쳐 적지 않는다.
    """
    gone = 0
    for table in document.tables:
        for row in table.rows:
            for cell in raw_cells(row):
                if has_diag(cell) and NA_ONLY.match(cell_text(cell) or ""):
                    set_cell(cell, "")
                    gone += 1
    return gone


def highlight(document, needle="확인 필요", color="yellow"):
    """그 글이 든 런에 형광펜을 칠한다. 칠한 런 수를 돌려준다.

    담당자 지시(2026-09): "확인 필요 내용은 노랑 마크를 칠해줘" — 채우지 못해 담당자가
    직접 봐야 하는 칸을 한눈에 찾게 한다.
    """
    n = 0
    for run in document.element.body.iter(qn("w:r")):
        text = "".join(t.text or "" for t in run.findall(qn("w:t")))
        if needle not in text:
            continue
        rpr = run.get_or_add_rPr()
        el = get_or_add(rpr, "highlight")
        el.set(qn("w:val"), color)
        n += 1
    return n


# ---------- 문서 전체 ----------
def loose(text):
    """빈칸 차이를 지운 꼴 — 같은 문구라도 결재본마다 일반 공백과 \xa0(줄바꿈 없는 공백)가 섞여 있다.

    디겐타안연고 결재본의 'QC-126\xa0제품품질평가규정' 을 퀴노비드 기준의 'QC-126 제품품질평가규정'
    으로 찾지 못해 엔진이 9항에서 멈췄다(2026-09). 한 글자도 다르지 않은데 빈칸 종류만 달랐다.
    """
    return re.sub(r"[\s\u00a0]+", " ", text or "").strip()


def find_para(document, needle):
    """문구로 문단을 찾는다 — 그대로 찾아보고, 없으면 빈칸 차이를 무시하고 다시 찾는다."""
    for p in document.paragraphs:
        if needle in p.text:
            return p
    key = loose(needle)
    if key:
        for p in document.paragraphs:
            if key in loose(p.text):
                return p
    return None


def replace_in_para(document, needle, old, new):
    p = find_para(document, needle)
    if p is None:
        raise KeyError(needle)
    if old in p.text:
        set_para_text(p, p.text.replace(old, new))
    else:                        # 빈칸이 \xa0 인 결재본 — 빈칸 종류를 가리지 않고 바꾼다
        pattern = "[\\s\\u00a0]+".join(re.escape(part) for part in re.split(r"[\s\u00a0]+", old))
        set_para_text(p, re.sub(pattern, lambda m: new, p.text))
    return p


def audit(path):
    """빈 페이지를 만드는 세 가지 함정 검사."""
    import zipfile
    from xml.etree import ElementTree as ET

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    body = root.find(W + "body")
    hard = len(root.findall(".//" + W + "br[@" + W + "type='page']"))
    incell = 0
    for tc in root.iter(W + "tc"):
        for pbb in tc.iter(W + "pageBreakBefore"):
            incell += 1
    floating = sum(1 for tbl in root.iter(W + "tbl")
                   if tbl.find(W + "tblPr/" + W + "tblpPr") is not None
                   or tbl.find(W + "tblPr/" + W + "tblOverlap") is not None)
    trailing = 0
    kids = list(body)
    for i, el in enumerate(kids[:-1]):
        if el.tag == W + "tbl":
            nxt = kids[i + 1]
            if nxt.tag == W + "p" and not "".join(t.text or "" for t in nxt.iter(W + "t")).strip():
                pr = nxt.find(W + "pPr")
                if pr is None or pr.find(W + "pageBreakBefore") is None:
                    trailing += 1
    return {"하드 페이지 나눔": hard, "셀 안 pageBreakBefore": incell,
            "떠 있는 표": floating, "표 뒤 빈 문단": trailing}


def raw_cells(row):
    """python-docx 의 vMerge 해석을 거치지 않고 그 행이 실제로 가진 셀만 돌려준다."""
    from docx.table import _Cell
    return [_Cell(tc, row.table) for tc in row._tr.findall(qn("w:tc"))]


def grid_width(table):
    grid = table._tbl.find(qn("w:tblGrid"))
    return len(grid.findall(qn("w:gridCol"))) if grid is not None else 0


def grid_cells(row, width=None):
    """그리드 열 번호 → 그 행의 칸. 가로 병합(gridSpan)은 첫 열에만 담는다.

    요약 행은 첫 칸이 두 열을 덮는 일이 많아(‘최댓값’ 이 연번·Lot No. 를 함께 덮는다),
    자리로 세면 값이 한 칸씩 밀린다 — 7항 수율에서 실제로 그렇게 나왔다.
    """
    width = grid_width(row.table) if width is None else width
    out, col = {}, 0
    for cell in raw_cells(row):
        pr = cell._tc.find(qn("w:tcPr"))
        span_el = pr.find(qn("w:gridSpan")) if pr is not None else None
        span = int(span_el.get(qn("w:val"))) if span_el is not None else 1
        if col < width:
            out[col] = cell
        col += span
    return out


def note_after(document, table, text, sample_needle="QC-126 제품품질평가규정"):
    """표 바로 뒤에 각주 문단을 새로 넣는다. 서식은 sample 문단에서 가져온다.

    본보기 문단이 없어도 각주는 넣는다 — 각주 하나 때문에 보고서 전체를 못 만들면 안 된다.
    """
    found = find_para(document, sample_needle)
    if found is None:
        new = etree.SubElement(table._tbl.getparent(), qn("w:p"))
        table._tbl.addnext(new)
        run = etree.SubElement(new, qn("w:r"))
        t = etree.SubElement(run, qn("w:t"))
        t.text = text
        t.set(qn("xml:space"), "preserve")
        return new
    sample = found._p
    new = copy.deepcopy(sample)
    for r in new.findall(qn("w:r")):
        new.remove(r)
    for b in new.findall(qn("w:bookmarkStart")) + new.findall(qn("w:bookmarkEnd")):
        new.remove(b)
    run = copy.deepcopy([r for r in sample.findall(qn("w:r")) if r.find(qn("w:t")) is not None][0])
    for t in run.findall(qn("w:t")):
        t.text = text
        t.set(qn("xml:space"), "preserve")
    new.append(run)
    table._tbl.addnext(new)
    return new


def set_cell_plain(cell, *lines):
    """set_cell 과 같되, 자동 번호(numPr)와 들여쓰기(ind)를 지운다."""
    set_cell(cell, *lines)
    for p in cell_paras(cell._tc):
        pr = p.find(qn("w:pPr"))
        if pr is None:
            continue
        for tag in ("w:numPr", "w:ind"):
            el = pr.find(qn(tag))
            if el is not None:
                pr.remove(el)


def page_break_before(para):
    p = getattr(para, "_p", para)
    pr = p.find(qn("w:pPr"))
    if pr is None:
        pr = p.makeelement(qn("w:pPr"), {})
        p.insert(0, pr)
    get_or_add(pr, "pageBreakBefore")


def is_blank_para(el):
    """글자·줄바꿈·그림·구역 나누기가 없는 '진짜 빈' 문단인지.

    Word 가 변환한 결재본은 표지 끝의 쪽 나눔이 문단 속성(sectPr)으로 들어 있어, 글자가 없다고
    지우면 표지와 결재표가 한 쪽에 붙어 버린다(담당자 PC 에서 실제로 생김). 그래서 구역 나누기·
    쪽 나눔·그리기 개체·'앞에서 쪽 나눔' 이 있는 문단은 빈 문단으로 보지 않는다.
    """
    if not el.tag.endswith("}p"):
        return False
    if "".join(t.text or "" for t in el.iter(qn("w:t"))).strip():
        return False
    if el.find(".//" + qn("w:br")) is not None:
        return False
    for tag in ("w:pict", "w:drawing", "w:sectPr", "w:pageBreakBefore", "w:object"):
        if el.find(".//" + qn(tag)) is not None:
            return False
    return True


def drop_blank_paras_between(document, table, para):
    """표와 문단 사이의 빈 문단(하드 나눔이 없는 것)을 지운다."""
    body = document.element.body
    kids = list(body)
    start = kids.index(table._tbl) + 1
    end = kids.index(getattr(para, "_p", para))
    removed = 0
    for el in kids[start:end]:
        if not is_blank_para(el):
            continue
        body.remove(el)
        removed += 1
    return removed


def drop_blank_paras_after(document, needle, count):
    """needle 문단 뒤에 이어지는 빈 문단을 count 개만 지운다."""
    body = document.element.body
    start = find_para(document, needle)._p
    kids = list(body)
    i = kids.index(start) + 1
    removed = 0
    while i < len(kids) and removed < count:
        el = kids[i]
        if not is_blank_para(el):
            break
        body.remove(el)
        removed += 1
        i += 1
    return removed


def set_line_spacing(para, line, rule="auto", before=None, after=None):
    p = getattr(para, "_p", para)
    pr = p.find(qn("w:pPr"))
    if pr is None:
        pr = p.makeelement(qn("w:pPr"), {})
        p.insert(0, pr)
    sp = get_or_add(pr, "spacing")
    sp.set(qn("w:line"), str(line))
    sp.set(qn("w:lineRule"), rule)
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))


def set_cell_align(cell, jc="center"):
    for p in cell_paras(cell._tc):
        pr = p.find(qn("w:pPr"))
        if pr is None:
            pr = p.makeelement(qn("w:pPr"), {})
            p.insert(0, pr)
        el = get_or_add(pr, "jc")
        el.set(qn("w:val"), jc)


def blank_para_before(document, needle, template_needle=None):
    """needle 문단 바로 앞에 빈 문단을 하나 넣는다."""
    target = find_para(document, needle)._p
    src = find_para(document, template_needle)._p if template_needle else target
    new = copy.deepcopy(src)
    for r in new.findall(qn("w:r")):
        new.remove(r)
    for b in new.findall(qn("w:bookmarkStart")) + new.findall(qn("w:bookmarkEnd")):
        new.remove(b)
    pr = new.find(qn("w:pPr"))
    if pr is not None:
        # 자동 번호(글머리 기호)를 물려받으면 빈 문단에 점만 찍혀 남는다.
        for tag in ("w:pageBreakBefore", "w:numPr"):
            got = pr.find(qn(tag))
            if got is not None:
                pr.remove(got)
    target.addprevious(new)
    return new


def set_cell_valign(cell, val="center"):
    pr = _tcpr(cell._tc)
    el = get_or_add(pr, "vAlign")
    el.set(qn("w:val"), val)


def hard_breaks_to_page_break_before(document):
    """'빈 문단 + 하드 페이지 나눔' 을 다음 제목의 pageBreakBefore 로 바꾼다.

    빈 문단이 페이지 맨 위에 홀로 남아 빈 쪽을 만드는 것을 막는다.
    다음이 표인 경우에는 표 앞 페이지 나눔이 필요하므로 그대로 둔다.
    """
    body = document.element.body
    moved = 0
    for el in list(body):
        if not el.tag.endswith("}p"):
            continue
        if el.find(".//" + qn("w:br")) is None:
            continue
        if "".join(t.text or "" for t in el.iter(qn("w:t"))).strip():
            continue                      # 글자가 있는 문단은 그대로
        nxt = el.getnext()
        if nxt is None or not nxt.tag.endswith("}p"):
            continue                      # 표 앞 페이지 나눔은 유지
        if not "".join(t.text or "" for t in nxt.iter(qn("w:t"))).strip():
            continue
        page_break_before(nxt)
        body.remove(el)
        moved += 1
    return moved


def has_drawing(table):
    return (table._tbl.find(".//" + qn("w:pict")) is not None
            or table._tbl.find(".//" + qn("w:drawing")) is not None)


def diag_empty_remarks(table, header="비고"):
    """머리행에서 header 열을 찾아, 비어 있는 (병합) 칸마다 사선을 하나 긋는다."""
    if has_drawing(table):
        return 0            # 그리기 개체 사선이 이미 있는 표는 건드리지 않는다
    try:
        cols = table.rows[0].cells
    except Exception:
        return 0
    import re as _re
    idx = [i for i, c in enumerate(cols)
           if _re.sub(r"\s+", "", cell_text(c)) == header]
    if not idx:
        return 0
    ci = idx[0]
    seen, n = set(), 0
    for row in table.rows:
        try:
            cell = row.cells[ci]
        except Exception:
            continue
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        if cell_text(cell).strip():
            continue
        add_diag(cell)
        n += 1
    return n


def _tc_at(tr, index):
    """그 행에서 그리드 index 열을 차지하는 칸. 가로 병합(gridSpan)을 세어 찾는다."""
    col = 0
    for tc in tr.findall(qn("w:tc")):
        pr = tc.find(qn("w:tcPr"))
        span_el = pr.find(qn("w:gridSpan")) if pr is not None else None
        span = int(span_el.get(qn("w:val"))) if span_el is not None else 1
        if col <= index < col + span:
            return tc
        col += span
    return None


def merge_empty_runs(table, header="비고"):
    """머리행에서 header 열을 찾아, **내용이 없는 연속 칸을 하나로 합치고** 사선을 하나 긋는다.

    담당자 지시(2026-09): 3항 대상 제품표의 비고 칸은 빈 줄끼리 이어 붙여 병합하고 사선 하나만
    긋는다 — 줄마다 'N/A' 를 적거나 사선을 여러 개 긋지 않는다.
    돌려주는 값: (합친 묶음 수, 합친 줄 수)
    """
    from docx.table import _Cell
    trs = table._tbl.findall(qn("w:tr"))
    if len(trs) < 2:
        return 0, 0
    # 머리글은 '비 고' 처럼 사이가 벌어져 있는 것이 많다 — 빈칸을 아예 지우고 견준다.
    squeeze = lambda text: re.sub(r"[\s\u00a0]+", "", text or "")
    want = squeeze(header)
    index = None
    col = 0
    for tc in trs[0].findall(qn("w:tc")):
        pr = tc.find(qn("w:tcPr"))
        span_el = pr.find(qn("w:gridSpan")) if pr is not None else None
        span = int(span_el.get(qn("w:val"))) if span_el is not None else 1
        if squeeze("".join(t.text or "" for t in tc.iter(qn("w:t")))) == want:
            index = col
            break
        col += span
    if index is None:
        return 0, 0

    cells = []
    for tr in trs[1:]:
        tc = _tc_at(tr, index)
        cells.append(tc)
    groups, run = [], []
    for tc in cells:
        empty = tc is not None and not "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
        if empty:
            run.append(tc)
        else:
            if run:
                groups.append(run)
            run = []
    if run:
        groups.append(run)

    merged_rows = 0
    for run in groups:
        for i, tc in enumerate(run):
            cell = _Cell(tc, table)
            set_vmerge(cell, "restart" if i == 0 else None)
            if i == 0:
                add_diag(cell)
            else:
                clear_diag(cell)
        merged_rows += len(run)
    return len(groups), merged_rows


def _keep_next(row):
    for tc in row._tr.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            pr = p.find(qn("w:pPr"))
            if pr is None:
                pr = p.makeelement(qn("w:pPr"), {})
                p.insert(0, pr)
            get_or_add(pr, "keepNext")


def keep_merged_groups(table, col):
    """col 열이 세로 병합된 묶음이 쪽 경계에서 갈라지지 않게 '다음과 함께' 를 건다."""
    rows = table.rows
    n = 0
    for i in range(len(rows) - 1):
        try:
            a = rows[i].cells[col]._tc
            b = rows[i + 1].cells[col]._tc
        except Exception:
            continue
        if a is b:
            _keep_next(rows[i])
            n += 1
    return n


def strip_blanks_before_page_breaks(document):
    """'앞에서 쪽 나눔' 이 걸린 문단 바로 앞의 빈 문단들을 지운다.

    Word 는 이 빈 문단들이 앞 쪽에 다 못 들어가면 다음 쪽으로 넘겼다가
    쪽 나눔을 만나 빈 쪽 하나를 통째로 만든다.
    """
    body = document.element.body
    removed = 0
    for el in list(body):
        if not el.tag.endswith("}p"):
            continue
        pr = el.find(qn("w:pPr"))
        if pr is None or pr.find(qn("w:pageBreakBefore")) is None:
            continue
        prev = el.getprevious()
        while prev is not None and prev.tag.endswith("}p"):
            if not is_blank_para(prev):
                break
            gone = prev
            prev = prev.getprevious()
            body.remove(gone)
            removed += 1
    return removed


def _has_page_break(el):
    return any(b.get(qn("w:type")) == "page" for b in el.iter(qn("w:br")))


def tidy_page_breaks(document):
    """빈 쪽이 생기는 구조를 없앤다. {'앞 빈 문단': n, '쪽나눔 전환': n, '끝 빈 문단': n}

    Word 는 LibreOffice 보다 줄이 높아 같은 문서가 몇 쪽 더 나온다(2026 퀴노비드: 32 → 37).
    그 차이가 쌓이면 '쪽나눔 문단 앞의 빈 문단' 이 다음 쪽으로 밀리고, 거기서 쪽나눔을 만나
    빈 쪽을 통째로 만든다. 그래서
      1) 하드 쪽나눔 문단 바로 앞의 빈 문단을 지운다 — 쪽 끝 채움일 뿐이라 배치가 안 바뀐다
      2) '빈 문단 + 하드 쪽나눔' 은 다음 문단의 pageBreakBefore 로 바꾼다 (표 앞이면 그대로)
      3) 문서 맨 끝의 빈 문단을 지운다 — 마지막 빈 쪽에 홀로 남던 것. 단, 표 바로 뒤의
         빈 문단은 둔다(Word 는 본문이 표로 끝나는 것을 허용하지 않는다).
    """
    body = document.element.body
    before = 0
    for el in list(body):
        if not el.tag.endswith("}p") or not _has_page_break(el):
            continue
        prev = el.getprevious()
        while prev is not None and prev.tag.endswith("}p") and is_blank_para(prev):
            gone, prev = prev, prev.getprevious()
            body.remove(gone)
            before += 1
    moved = hard_breaks_to_page_break_before(document)
    before += strip_blanks_before_page_breaks(document)
    tail = 0
    for el in reversed(list(body)):
        if el.tag.endswith("}sectPr"):
            continue
        prev = el.getprevious()
        if (el.tag.endswith("}p") and is_blank_para(el)
                and prev is not None and not prev.tag.endswith("}tbl")):
            body.remove(el)
            tail += 1
        else:
            break
    return {"앞 빈 문단": before, "쪽나눔 전환": moved, "끝 빈 문단": tail}


def fix_table_widths(table):
    """표 배치를 '고정' 으로 두고 칸 폭을 그리드(w:tblGrid)와 맞춘다. 고친 칸 수를 돌려준다.

    EDMS 서식은 칸 폭(w:tcW)이 행마다 단위가 뒤섞여 있다 — 머리행은 pct 5000 기준, 요약
    행은 dxa 값을 pct 로 적어 3932% 같은 값이 들어 있다. Word 는 자동 맞춤에서 이 값을 함께
    보므로 그리드와 다르게 열을 나눈다(성상은 넓고 생균수는 좁아 결과가 두 줄, 튜브인쇄는
    네 줄). 담당자 지적: "되도록 한 줄에 결과가 표시되도록 오른쪽처럼", "한 줄에 안 되면
    두 줄로". 그리드는 담당자가 의도한 폭이므로 그리드를 그대로 쓰게 한다.
    """
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return 0
    cols = [int(g.get(qn("w:w")) or 0) for g in grid.findall(qn("w:gridCol"))]
    if not cols or not all(cols):
        return 0
    pr = get_or_add(tbl, "tblPr")
    layout = get_or_add(pr, "tblLayout")
    layout.set(qn("w:type"), "fixed")
    tw = get_or_add(pr, "tblW")
    tw.set(qn("w:w"), str(sum(cols)))
    tw.set(qn("w:type"), "dxa")
    n = 0
    for tr in tbl.findall(qn("w:tr")):
        col = 0
        for tc in tr.findall(qn("w:tc")):
            tcpr = _tcpr(tc)
            span_el = tcpr.find(qn("w:gridSpan"))
            span = int(span_el.get(qn("w:val"))) if span_el is not None else 1
            width = sum(cols[col:col + span]) or cols[-1]
            w = get_or_add(tcpr, "tcW")
            w.set(qn("w:w"), str(width))
            w.set(qn("w:type"), "dxa")
            col += span
            n += 1
    return n


def text_width(document, section=0):
    """본문 폭(dxa) — 쪽 폭에서 좌우 여백을 뺀 값."""
    sect = document.sections[section]
    pg = sect._sectPr.find(qn("w:pgSz"))
    mar = sect._sectPr.find(qn("w:pgMar"))
    if pg is None or mar is None:
        return 0
    try:
        return int(pg.get(qn("w:w"))) - int(mar.get(qn("w:left"))) - int(mar.get(qn("w:right")))
    except (TypeError, ValueError):
        return 0


def fit_to_window(table, width):
    """표를 '창에 자동으로 맞춤' 으로 — 본문 폭에 딱 맞춘다. 바꿨으면 True.

    담당자 지시(2026-09): "표 크기는 창에 자동으로 맞춤으로 조정해줘. 모든 표에 해당이 돼."
    전년도 결재본에서 물려받은 표는 본문 폭(9978)보다 넓은 것이 있어(10105) 오른쪽 여백을
    넘었다. 그리드를 본문 폭에 비례로 다시 나누고, 표 너비를 100%(pct 5000)로 둔다 —
    2026 결재본도 pct 로 되어 있다.
    """
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None or width <= 0:
        return False
    gcols = grid.findall(qn("w:gridCol"))
    cols = [int(g.get(qn("w:w")) or 0) for g in gcols]
    if not cols or not all(cols):
        return False
    total = sum(cols)
    new = [max(1, int(round(c * width / float(total)))) for c in cols]
    new[-1] = max(1, new[-1] + width - sum(new))          # 반올림 오차는 마지막 열에서 맞춘다
    for g, w in zip(gcols, new):
        g.set(qn("w:w"), str(w))
    for tr in tbl.findall(qn("w:tr")):
        col = 0
        for tc in tr.findall(qn("w:tc")):
            tcpr = _tcpr(tc)
            span_el = tcpr.find(qn("w:gridSpan"))
            span = int(span_el.get(qn("w:val"))) if span_el is not None else 1
            w = get_or_add(tcpr, "tcW")
            w.set(qn("w:w"), str(sum(new[col:col + span]) or new[-1]))
            w.set(qn("w:type"), "dxa")
            col += span
    pr = get_or_add(tbl, "tblPr")
    tw = get_or_add(pr, "tblW")
    tw.set(qn("w:w"), "5000")
    tw.set(qn("w:type"), "pct")
    get_or_add(pr, "tblLayout").set(qn("w:type"), "autofit")
    ind = get_or_add(pr, "tblInd")
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    return True


UNIT = 92           # 반각 한 자의 폭(twip) — 굴림 10pt 를 Word 로 재어 맞춤. 한글·전각은 2 단위
CELL_MARGIN = 220   # 칸 좌우 여백(108×2) + 여유
KEY_HEADS = ("연번", "no.", "lot no", "lot", "구분")
NOWRAP_HEADS = ("번호", "no.", "lot")     # 제조 번호·문서 번호 — 줄을 넘기면 읽을 수 없다


def text_units(text):
    """글의 시각적 길이(반각 단위). 한글·전각 2, 나머지 1. 줄바꿈이 있으면 가장 긴 줄."""
    best = 0
    for line in (text or "").split("\n"):
        n = 0
        for ch in line:
            n += 2 if ord(ch) > 0x2E7F else 1
        best = max(best, n)
    return best


def balance_columns(table, fixed_heads=KEY_HEADS, unit=UNIT, margin=CELL_MARGIN):
    """열마다 글 길이를 재서 줄 수가 최대한 같아지도록 그리드 폭을 다시 나눈다. 새 그리드를 돌려준다.

    담당자 지적: "오른편처럼 최대한 균등하게" — 성상 열이 넓어 한 줄이고 확인 2) 열이 좁아
    아홉 줄이면 안 된다. 연번·Lot No. 같은 열쇠 열은 그대로 두고, 나머지 폭을
      1) 모든 열이 L 줄 안에 들어가는 가장 작은 L 을 찾아 그만큼 배정하고
      2) 남는 폭은 열마다 똑같이 나눈다
    그래서 짧은 값(10 CFU/g 미만·음성)은 한 줄, 긴 문장끼리는 같은 줄 수가 된다.
    가로 병합(gridSpan) 칸은 한 열의 길이가 아니므로 재지 않는다. 표 배치는 fix_table_widths 로
    고정한다(이 함수가 마지막에 부른다).
    """
    tbl = table._tbl
    grid_el = tbl.find(qn("w:tblGrid"))
    if grid_el is None:
        return None
    cols = [int(g.get(qn("w:w")) or 0) for g in grid_el.findall(qn("w:gridCol"))]
    if len(cols) < 3 or not all(cols):
        return None
    need = [0] * len(cols)
    heads = [""] * len(cols)
    rows = tbl.findall(qn("w:tr"))
    for ri, tr in enumerate(rows):
        tcs = tr.findall(qn("w:tc"))
        first = "".join(t.text or "" for t in tcs[0].iter(qn("w:t"))).strip() if tcs else ""
        # 머리행은 담당자가 일부러 두 줄로 둔 제목이 많아 줄 수에 세지 않는다. 자료 행(연번이
        # 숫자인 행)만 잰다 — 요약 행의 값은 짧은 숫자라 어차피 한 줄이다.
        is_data = first.isdigit()
        col = 0
        for tc in tcs:
            pr = tc.find(qn("w:tcPr"))
            sp = pr.find(qn("w:gridSpan")) if pr is not None else None
            span = int(sp.get(qn("w:val"))) if sp is not None else 1
            if span == 1 and col < len(cols):
                txt = "".join(t.text or "" for t in tc.iter(qn("w:t")))
                if ri == 0:
                    heads[col] = txt.strip().lower()
                if is_data:
                    need[col] = max(need[col], text_units(txt))
            col += span
    if not any(need):
        return None
    fixed = [any(k in h for k in fixed_heads) for h in heads]
    free = [i for i in range(len(cols)) if not fixed[i]]
    if len(free) < 2:
        return None
    total = sum(cols)
    budget = total - sum(cols[i] for i in range(len(cols)) if fixed[i])
    floor = {i: 0.6 * cols[i] for i in free}          # 원래 폭의 60% 밑으로는 좁히지 않는다
    for i in free:                                    # 번호 칸은 줄을 넘기지 않는다 (OGW701 → OGW / 701)
        if any(k in heads[i] for k in NOWRAP_HEADS) and need[i]:
            floor[i] = max(floor[i], need[i] * unit + margin)
    chosen = best = None
    for lines in range(1, 40):
        want = [max(need[i] * unit / float(lines) + margin, floor[i]) for i in free]
        if sum(want) <= budget:
            chosen, best = want, lines
            break
    if chosen is None:
        return None
    # 지금 폭으로도 줄 수가 이미 최소면 담당자가 정한 폭을 그대로 둔다
    # (담당자: 함량·입자도·무균·질량 표는 "이것은 잘했어" — 다시 나누지 않는다).
    current = max(int(math.ceil(need[i] * unit / max(1.0, cols[i] - margin))) if need[i] else 1
                  for i in free)
    if current <= best:
        return None
    # 남는 폭은 열마다 똑같이 나눈다. 다만 빈 열(13.1 '실시 사유')이 있으면 글 길이에 비례해
    # 나눈다 — 빈 열에 같은 몫을 주면 'OGW70 / 1' 처럼 값이 든 열이 잘린다.
    left = budget - sum(chosen)
    empty = any(need[i] == 0 for i in free)
    weight = [float(need[i]) for i in free] if empty else [1.0] * len(free)
    if not sum(weight):
        weight = [1.0] * len(free)
    new = list(cols)
    for k, i in enumerate(free):
        new[i] = int(round(chosen[k] + left * weight[k] / sum(weight)))
    new[free[-1]] += total - sum(new)          # 반올림 오차는 마지막 자유 열에
    for g, w in zip(grid_el.findall(qn("w:gridCol")), new):
        g.set(qn("w:w"), str(w))
    fix_table_widths(table)
    return new


def fix_all_table_widths(document, skip=()):
    """본문 표 전부의 폭을 그리드에 맞춘다(skip 번호는 제외). 고친 칸 수."""
    n = 0
    for ti, table in enumerate(document.tables):
        if ti in skip:
            continue
        n += fix_table_widths(table)
    return n


def zero_cell_spacing(document, skip=()):
    """표 안 모든 문단의 줄 모양을 하나로 맞춘다 — 앞·뒤 간격 0, 줄 간격 1줄, 들여쓰기 없음.

    담당자 지시(2026-09): "모든 표 안의 모든 글씨는 동일하게 맞춰줘." 전년도 결재본은 칸마다
    줄 간격(240 exact · 315 atLeast · 360 auto …)과 들여쓰기(firstLine 600 · hanging 600 · 없음)가
    제각각이라, 같은 글인데도 조제 성상은 두 줄 사이가 벌어지고 충전 성상은 붙어 보였다.
    Style12 가 물려주는 앞뒤 14pt 도 함께 지운다(Word 의 '단락 앞/뒤 공백 제거').
    """
    n = 0
    skip_tbls = {document.tables[i]._tbl for i in skip}
    for tbl in document.element.body.iter(qn("w:tbl")):
        if tbl in skip_tbls:
            continue
        for tc in tbl.iter(qn("w:tc")):
            for p in tc.findall(qn("w:p")):
                pr = p.find(qn("w:pPr"))
                if pr is None:
                    pr = p.makeelement(qn("w:pPr"), {})
                    p.insert(0, pr)
                sp = get_or_add(pr, "spacing")
                sp.set(qn("w:before"), "0")
                sp.set(qn("w:after"), "0")
                sp.set(qn("w:line"), "240")            # 1줄 — 칸마다 다르던 줄 높이를 맞춘다
                sp.set(qn("w:lineRule"), "auto")
                for k in ("w:beforeLines", "w:afterLines", "w:beforeAutospacing", "w:afterAutospacing"):
                    if sp.get(qn(k)) is not None:
                        del sp.attrib[qn(k)]
                ind = pr.find(qn("w:ind"))             # 첫 줄·내어쓰기 들여쓰기는 칸마다 달라 지운다
                if ind is not None:
                    pr.remove(ind)
                n += 1
    return n


def hanging_indent_notes(document, width=480):
    """본문(표 밖)의 각주 문단 '1) …' 에 내어쓰기를 준다.

    번호 뒤에 탭을 넣고 둘째 줄부터는 번호 다음 글자 위치에 맞춘다
    (Word: 왼쪽 width, 첫 줄 -width).
    """
    import re as _re
    body = document.element.body
    n = 0
    for p in body.findall(qn("w:p")):          # 표 안 문단은 제외(직계 자식만)
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        m = _re.match(r"\s*(\d\))\s+", text)
        if not m:
            continue
        pr = p.find(qn("w:pPr"))
        if pr is None:
            pr = p.makeelement(qn("w:pPr"), {})
            p.insert(0, pr)
        ind = get_or_add(pr, "ind")
        ind.set(qn("w:left"), str(width))
        ind.set(qn("w:hanging"), str(width))
        # 번호와 본문 사이를 탭으로: 첫 w:t 를 '1)' / tab / 나머지 로 나눈다
        for t in p.iter(qn("w:t")):
            if t.text and _re.match(r"\s*\d\)\s+", t.text):
                head = m.group(1)
                rest = t.text[m.end():]
                t.text = head
                tab = t.makeelement(qn("w:tab"), {})
                t.addnext(tab)
                t2 = t.makeelement(qn("w:t"), {})
                t2.text = rest
                t2.set(qn("xml:space"), "preserve")
                tab.addnext(t2)
                break
        n += 1
    return n


def _span(tc):
    pr = tc.find(qn("w:tcPr"))
    g = pr.find(qn("w:gridSpan")) if pr is not None else None
    return int(g.get(qn("w:val"))) if g is not None else 1


def _set_span(tc, n):
    pr = _tcpr(tc)
    g = pr.find(qn("w:gridSpan"))
    if n <= 1:
        if g is not None:
            pr.remove(g)
        return
    if g is None:
        g = pr.makeelement(qn("w:gridSpan"), {})
        w = pr.find(qn("w:tcW"))
        (w.addnext(g) if w is not None else pr.insert(0, g))
    g.set(qn("w:val"), str(n))


def _width(tc):
    pr = tc.find(qn("w:tcPr"))
    w = pr.find(qn("w:tcW")) if pr is not None else None
    try:
        return int(w.get(qn("w:w")))
    except Exception:
        return 0


def slash_empty_summary(table, n_summary=5, rows=None):
    """9.2 세부표의 요약 5행(최댓값~Cpk 판정)에서 다섯 행 모두 비어 있는 값 열을
    가로·세로로 한 덩어리로 병합하고 사선(왼쪽 아래→오른쪽 위) 하나를 긋는다.
    (담당자 지시 — 정성 항목의 요약 칸 처리 방식)"""
    from docx.table import _Cell
    if rows is None:
        rows = table.rows[-n_summary:]
    if not rows or not cell_text(_Cell(rows[0]._tr.findall(qn("w:tc"))[0], table)).strip().startswith("최댓값"):
        return 0
    tcs = [r._tr.findall(qn("w:tc")) for r in rows]
    n = len(tcs[0])
    if any(len(t) != n for t in tcs) or n < 2:
        return 0
    empty = [all(not "".join(x.text or "" for x in tcs[r][c].iter(qn("w:t"))).strip()
                 for r in range(n_summary)) for c in range(n)]
    empty[0] = False                                   # 라벨 칸은 제외
    runs, c = [], 1
    while c < n:
        if empty[c]:
            a = c
            while c + 1 < n and empty[c + 1]:
                c += 1
            runs.append((a, c))
        c += 1
    made = 0
    for a, b in reversed(runs):                        # 뒤에서부터 지워야 앞 인덱스가 유지됨
        for r in range(n_summary):
            first = tcs[r][a]
            span = sum(_span(tcs[r][k]) for k in range(a, b + 1))
            width = sum(_width(tcs[r][k]) for k in range(a, b + 1))
            for k in range(b, a, -1):
                tcs[r][k].getparent().remove(tcs[r][k])
            _set_span(first, span)
            pr = _tcpr(first)
            w = pr.find(qn("w:tcW"))
            if w is not None and width:
                w.set(qn("w:w"), str(width))
            cell = _Cell(first, table)
            set_vmerge(cell, "restart" if r == 0 else None)
            clear_diag(cell)
        add_diag(_Cell(tcs[0][a], table))
        made += 1
    return made


def keep_headings_with_next(document):
    """본문 항 제목('7. 수율 현황', '7.1. 내수용' 등)에 '다음과 함께'를 걸어
    제목만 쪽 맨 아래에 남지 않게 한다."""
    import re as _re
    n = 0
    for p in document.element.body.findall(qn("w:p")):
        txt = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        if not _re.match(r"^\d{1,2}(\.\d+)*\.\s*\S", txt) or len(txt) > 60:
            continue
        pr = p.find(qn("w:pPr"))
        if pr is None:
            pr = p.makeelement(qn("w:pPr"), {})
            p.insert(0, pr)
        if pr.find(qn("w:keepNext")) is None:
            get_or_add(pr, "keepNext")
            n += 1
    return n


def superscript_note_marks(document):
    """표 안 값 끝에 붙은 각주 번호('82.621)' 의 '1)')를 윗첨자로 만든다."""
    import copy as _copy
    import re as _re
    changed = []
    for tbl in document.element.body.iter(qn("w:tbl")):
        for tc in tbl.iter(qn("w:tc")):
            for p in tc.findall(qn("w:p")):
                ts = [t for t in p.iter(qn("w:t")) if t.text]
                if not ts:
                    continue
                whole = "".join(t.text for t in ts)
                m = _re.search(r"(\d\))\s*$", whole)
                if not m or whole.strip() == m.group(1):
                    continue
                head = whole[: m.start()]
                if head.count("(") != head.count(")"):
                    continue          # '(rev.7)' '(MTC001)' 처럼 괄호 안 숫자는 각주가 아니다
                last = ts[-1]
                mm = _re.search(r"(\d\))\s*$", last.text)
                if not mm:
                    continue
                mark = mm.group(1)
                last.text = last.text[: mm.start()]
                run = last.getparent()
                new = _copy.deepcopy(run)
                for t in list(new.iter(qn("w:t"))):
                    t.getparent().remove(t)
                rpr = new.find(qn("w:rPr"))
                if rpr is None:
                    rpr = new.makeelement(qn("w:rPr"), {})
                    new.insert(0, rpr)
                va = get_or_add(rpr, "vertAlign")
                va.set(qn("w:val"), "superscript")
                t2 = new.makeelement(qn("w:t"), {})
                t2.text = mark
                t2.set(qn("xml:space"), "preserve")
                new.append(t2)
                run.addnext(new)
                changed.append(whole.strip()[:26])
    return changed


def add_diag_tl2br(cell):
    """왼쪽 위 → 오른쪽 아래 사선 (머리 칸을 둘로 가를 때)."""
    pr = _tcpr(cell._tc)
    borders = get_or_add(pr, "tcBorders")
    for tag in ("w:tl2br", "w:tr2bl"):
        el = borders.find(qn(tag))
        if el is not None:
            borders.remove(el)
    el = get_or_add(borders, "tl2br")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "4")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")


def split_header_cell(cell, top_right, bottom_left):
    """'기준 / Lot No.' 처럼 머리 칸을 사선으로 갈라, 위 글자는 오른쪽,
    아래 글자는 왼쪽으로 붙인다."""
    paras = cell_paras(cell._tc)
    for i, p in enumerate(paras):
        txt = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        pr = p.find(qn("w:pPr"))
        if pr is None:
            pr = p.makeelement(qn("w:pPr"), {})
            p.insert(0, pr)
        old = pr.find(qn("w:jc"))
        if old is not None:
            pr.remove(old)
        jc = pr.makeelement(qn("w:jc"), {})
        jc.set(qn("w:val"), "right" if txt == top_right else
               ("left" if txt == bottom_left else "center"))
        pr.append(jc)
    add_diag_tl2br(cell)


def diag_empty_in_column(table, header):
    """머리행에서 header 열을 찾아 비어 있는 칸마다 사선(↗)을 긋는다."""
    import re as _re
    try:
        cols = table.rows[0].cells
    except Exception:
        return 0
    idx = [i for i, c in enumerate(cols) if _re.sub(r"\s+", "", cell_text(c)) == header]
    if not idx:
        return 0
    ci, seen, n = idx[0], set(), 0
    for row in table.rows:
        try:
            cell = row.cells[ci]
        except Exception:
            continue
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        if cell_text(cell).strip():
            continue
        add_diag(cell)
        n += 1
    return n


def table_font_size(document, half_points=20, keep_smaller=True):
    """표 안 글씨를 굴림 10(sz 20)으로 맞춘다.
    keep_smaller 이면 이미 10보다 작게 줄여 둔 칸(일탈·변경 서술)은 그대로 둔다."""
    n = 0
    for tbl in document.element.body.iter(qn("w:tbl")):
        for r in tbl.iter(qn("w:r")):
            if not any(t.text for t in r.findall(qn("w:t"))):
                continue
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                rpr = r.makeelement(qn("w:rPr"), {})
                r.insert(0, rpr)
            for tag in ("w:sz", "w:szCs"):
                el = rpr.find(qn(tag))
                if el is not None and keep_smaller:
                    try:
                        if int(el.get(qn("w:val"))) < half_points:
                            continue
                    except (TypeError, ValueError):
                        pass
                if el is None:
                    el = get_or_add(rpr, tag.split(":")[1])
                el.set(qn("w:val"), str(half_points))
            n += 1
    return n


def keep_table_together(table):
    """표 전체를 한 쪽에 붙여 둔다. 한 쪽에 다 못 들어가면 통째로 다음 쪽에서 시작한다."""
    rows = table.rows
    for r in rows[:-1]:
        _keep_next(r)
    return len(rows) - 1


def trim_row_heights(document, floor=340, skip=()):
    """행의 '최소 높이'가 필요 이상으로 크면 floor 로 낮춘다.

    hRule=atLeast 는 '최소'라서 낮춰도 글이 잘리지 않는다(내용이 길면 Word 가
    알아서 늘린다). 표가 한 줄 때문에 다음 쪽으로 넘어가는 것을 막는다.
    """
    n = 0
    skip_tbls = {document.tables[i]._tbl for i in skip}
    for tbl in document.element.body.iter(qn("w:tbl")):
        if tbl in skip_tbls:
            continue
        for tr in tbl.iter(qn("w:tr")):
            pr = tr.find(qn("w:trPr"))
            if pr is None:
                continue
            h = pr.find(qn("w:trHeight"))
            if h is None:
                continue
            if h.get(qn("w:hRule")) not in (None, "atLeast"):
                continue
            try:
                v = int(h.get(qn("w:val")))
            except (TypeError, ValueError):
                continue
            if v > floor:
                h.set(qn("w:val"), str(floor))
                n += 1
    return n


NO_BORDER = ("nil", "none")


def _has_diag(tc):
    """칸에 사선이 실제로 그어져 있는지 본다.

    Word 는 '사선 없음'도 <w:tr2bl w:val="nil"/> 로 적어 둔다. 요소가 있는지만 보면
    사선이 없는 칸을 있다고 잘못 읽으므로 w:val 까지 확인한다.
    """
    pr = tc.find(qn("w:tcPr"))
    b = pr.find(qn("w:tcBorders")) if pr is not None else None
    if b is None:
        return False
    for tag in ("w:tr2bl", "w:tl2br"):
        e = b.find(qn(tag))
        if e is not None and (e.get(qn("w:val")) or "single") not in NO_BORDER:
            return True
    return False


def diag_all_empty(document, skip=()):
    """GMP 문서는 공란을 남기지 않는다 — 표의 빈 칸마다 사선(↗)을 하나 긋는다.

    세로 병합으로 이어진 칸(위 칸의 연장)과 이미 사선이 있는 칸은 건드리지 않는다.
    skip 에는 결재표·개정 내역처럼 빈칸을 그대로 두는 표의 번호를 넘긴다.
    """
    from docx.table import _Cell
    n = 0
    for ti, tbl in enumerate(document.tables):
        if ti in skip:
            continue
        for tr in tbl._tbl.findall(qn("w:tr")):
            for tc in tr.findall(qn("w:tc")):
                pr = tc.find(qn("w:tcPr"))
                vm = pr.find(qn("w:vMerge")) if pr is not None else None
                if vm is not None and vm.get(qn("w:val")) != "restart":
                    continue                      # 병합으로 이어진 칸
                if _has_diag(tc):
                    continue
                if "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip():
                    continue
                add_diag(_Cell(tc, tbl))
                n += 1
    return n


def no_split_rows(document):
    """행 하나가 쪽 경계에서 반으로 잘리지 않게 한다(Word: '행 분할 허용' 해제).

    표 자체는 행과 행 사이에서 나뉠 수 있으므로, 긴 표도 빈 공간 없이 이어진다.
    """
    n = 0
    for tbl in document.element.body.iter(qn("w:tbl")):
        for tr in tbl.iter(qn("w:tr")):
            pr = tr.find(qn("w:trPr"))
            if pr is None:
                pr = tr.makeelement(qn("w:trPr"), {})
                tr.insert(0, pr)
            if pr.find(qn("w:cantSplit")) is None:
                get_or_add(pr, "cantSplit")
                n += 1
    return n


def table_chars(table):
    return sum(len("".join(t.text or "" for t in tc.iter(qn("w:t"))))
               for tc in table._tbl.iter(qn("w:tc")))


# 일탈·변경 서술 칸의 들여쓰기 — 기준 본(결재본)에서 쓰던 값 그대로
IND_KINDS = {
    "none": {},
    "body": {"hanging": "218", "start": "247", "end": "0"},   # '- ' 뒤 내어쓰기
    "cont": {"start": "247", "end": "0"},                     # 이어지는 문단
    "num":  {"firstLine": "180", "end": "0"},                 # '1) …' 줄
}


def set_cell_flow(cell, items):
    """(글, 들여쓰기 종류) 목록으로 칸을 채운다.

    줄바꿈을 손으로 넣지 않고 Word 가 칸 너비에 맞춰 접게 두므로,
    글꼴·칸 너비가 달라져도 줄이 어색하게 잘리지 않는다.
    """
    set_cell_plain(cell, *[t for t, _ in items])
    paras = cell_paras(cell._tc)
    for p, (_, kind) in zip(paras, items):
        spec = IND_KINDS[kind]
        pr = p.find(qn("w:pPr"))
        if pr is None:
            pr = p.makeelement(qn("w:pPr"), {})
            p.insert(0, pr)
        old = pr.find(qn("w:ind"))
        if old is not None:
            pr.remove(old)
        if not spec:
            continue
        ind = get_or_add(pr, "ind")
        for k, v in spec.items():
            ind.set(qn("w:" + k), v)


def slash_block(table, first, last):
    """first..last 행의 모든 칸을 한 칸으로 병합하고 사선(↗) 하나를 긋는다.

    '해당 없음' 표(8.2.3·11.1.2·11.2·14항·15항)의 빈 블록용. 결재본은 위치가 고정된
    그리기 개체 선을 썼는데, 행 높이가 바뀌면 선이 표 밖으로 나가므로 칸에 고정되는
    셀 테두리 사선으로 바꾼다. (담당자: "칸이 변경되면 사선도 위치에 고정")
    """
    from docx.table import _Cell
    rows = table._tbl.findall(qn("w:tr"))[first:last + 1]
    for r, tr in enumerate(rows):
        tcs = tr.findall(qn("w:tc"))
        head = tcs[0]
        span = sum(_span(tc) for tc in tcs)
        width = sum(_width(tc) for tc in tcs)
        for tc in tcs[1:]:
            tr.remove(tc)
        _set_span(head, span)
        pr = _tcpr(head)
        w = pr.find(qn("w:tcW"))
        if w is not None and width:
            w.set(qn("w:w"), str(width))
        cell = _Cell(head, table)
        set_vmerge(cell, "restart" if r == 0 else None)
        clear_diag(cell)
        for p in head.findall(qn("w:p")):          # 빈 칸이어야 한다 — 남은 글자 제거
            for t in p.iter(qn("w:t")):
                t.text = ""
    add_diag(_Cell(rows[0].findall(qn("w:tc"))[0], table))
    return 1


def strip_floating_lines(document, anchors):
    """anchors(본문 요소 목록) 안의 그리기 개체(pict/drawing/AlternateContent) 런을 지운다."""
    MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
    n = 0
    for el in anchors:
        for r in list(el.iter(qn("w:r"))):
            has = (r.find(".//" + qn("w:pict")) is not None
                   or r.find(".//" + qn("w:drawing")) is not None
                   or r.find(".//" + MC) is not None)
            if not has:
                continue
            if "".join(t.text or "" for t in r.iter(qn("w:t"))).strip():
                continue
            r.getparent().remove(r)
            n += 1
        for ac in list(el.iter(MC)):                # 런 밖에 남은 AlternateContent
            ac.getparent().remove(ac)
            n += 1
    return n


NOTE_START = re.compile(r"^\s*(?:[*※•]|\d{1,2}\))")


def drop_blank_after_note(document):
    """각주 줄과 다음 항 제목 사이의 빈 줄을 없앤다. 지운 수를 돌려준다.

    각주('* [RSF102] 플루오로메톨론 사용이력 없음.')는 바로 위 표에 딸린 글이라, 그 아래
    빈 줄까지 두면 다음 항이 한 줄씩 밀려 표가 다음 쪽으로 넘어간다
    (담당자: "쓸데없는 엔터는 지워줘").
    """
    body = document.element.body
    kids = list(body)

    def txt(el):
        return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()

    removed = 0
    for i, el in enumerate(kids):
        if not is_blank_para(el) or i == 0 or i + 1 >= len(kids):
            continue
        before, after = kids[i - 1], kids[i + 1]
        if before.tag != qn("w:p") or after.tag != qn("w:p"):
            continue
        if NOTE_START.match(txt(before)) and re.match(r"^\d{1,2}(\.\d+)*\.?\s*\S", txt(after)):
            body.remove(el)
            removed += 1
    return removed


def collapse_blank_runs(document, keep=1, min_run=2, protect_before=()):
    """본문에서 빈 문단이 min_run 개 이상 이어지면 keep 개만 남긴다.

    결재본은 9.2 표 사이에 빈 문단 4~8개를 넣어 다음 표를 다음 쪽으로 밀었는데,
    내용이 바뀌면 그 빈 줄이 쪽 중간의 공백으로 남아 표가 쪽 끝에서 갈라진다
    (담당자: "표 위치를 조절해야지 쓸데없이 공백이 있잖아"). 표 배치는 keepNext 로 하고
    빈 줄은 한 줄만 둔다. protect_before 에 든 글로 시작하는 문단 앞 묶음은 건드리지 않는다.
    """
    body = document.element.body
    kids = list(body)

    def txt(el):
        return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()

    def blank(el):
        return is_blank_para(el)

    removed, i = 0, 0
    while i < len(kids):
        if not blank(kids[i]):
            i += 1
            continue
        j = i
        while j < len(kids) and blank(kids[j]):
            j += 1
        nxt = kids[j] if j < len(kids) else None
        nxt_txt = txt(nxt) if nxt is not None else ""
        if j - i >= min_run and not any(nxt_txt.startswith(k) for k in protect_before):
            for el in kids[i + keep:j]:
                body.remove(el)
                removed += 1
        i = j
    return removed


def keep_tail_together(table, n=6):
    """표의 마지막 n 행(요약 5행 + 그 위 데이터 1행)이 쪽 경계에서 갈라지지 않게 한다."""
    rows = table.rows
    for r in rows[-n:-1]:
        _keep_next(r)
    return min(n, len(rows)) - 1


def est_height(table, line=240):
    """표 높이 어림(dxa): 행마다 max(최소 높이, 가장 긴 칸의 줄 수 × 줄 높이)."""
    total = 0
    for tr in table._tbl.findall(qn("w:tr")):
        pr = tr.find(qn("w:trPr"))
        h = pr.find(qn("w:trHeight")) if pr is not None else None
        try:
            hv = int(h.get(qn("w:val")))
        except (AttributeError, TypeError, ValueError):
            hv = 0
        lines = max(len(tc.findall(qn("w:p"))) for tc in tr.findall(qn("w:tc")))
        total += max(hv, lines * line + 60)
    return total


def merged_groups(table, col):
    """col 열의 세로 병합 묶음을 [(first_row, last_row), ...] 로 돌려준다."""
    rows = table.rows
    groups, start = [], None
    for i in range(len(rows)):
        try:
            tc = rows[i].cells[col]._tc
        except Exception:
            if start is not None:
                groups.append((start, i - 1)); start = None
            continue
        if start is None:
            start = i
        elif tc is not rows[start].cells[col]._tc:
            groups.append((start, i - 1)); start = i
    if start is not None:
        groups.append((start, len(rows) - 1))
    return [g for g in groups if g[1] > g[0]]


def keep_groups(table, col, max_rows):
    """col 열 묶음이 max_rows 이하이면 묶음 전체에 keepNext (한 덩어리로 넘어가게).
    한 쪽에 안 들어갈 만큼 큰 묶음은 두지 않는다 — Word 는 쪽을 넘는 keepNext 사슬을
    통째로 무시하고 아무 데서나 자른다."""
    n = 0
    for a, b in merged_groups(table, col):
        if b - a + 1 > max_rows:
            continue
        for i in range(a, b):
            _keep_next(table.rows[i]); n += 1
    return n


def split_year_by_lot(table, year_col, lot_col, max_rows):
    """연도 묶음이 max_rows 를 넘으면 그 묶음의 연도 병합을 Lot 묶음 단위로 끊어,
    어느 쪽에서 시작하든 연도 글자가 보이게 한다."""
    from docx.table import _Cell
    n = 0
    lots = merged_groups(table, lot_col)
    for a, b in merged_groups(table, year_col):
        if b - a + 1 <= max_rows:
            continue
        label = cell_text(table.rows[a].cells[year_col]).strip()
        for la, lb in lots:
            if la < a or lb > b or la == a:
                continue
            first = raw_cells(table.rows[la])[year_col]
            set_vmerge(first, "restart")
            set_cell(first, label)
            n += 1
    return n


_VML_ID = [1000]


def collapse_empty_block(table, first, last):
    """비어 있는 first..last 행을 한 줄만 남기고 지운다. 지운 줄 수를 돌려준다.

    담당자 지시(2026-09): "내용이 없으면 1행만 남겨두고 사선처리해줘." 2026 결재본의 8.2.3
    신규 납품 제조원도 머리행 + 빈 줄 하나뿐이다. 빈 줄 셋에 사선을 걸치면 자리만 차지한다.
    """
    trs = table._tbl.findall(qn("w:tr"))
    gone = 0
    for tr in trs[first + 1:last + 1]:
        tr.getparent().remove(tr)
        gone += 1
    return gone


def draw_block_line(table, first, last, weight="0.5pt"):
    """first..last 행(빈 블록) 전체를 왼쪽 아래→오른쪽 위로 가로지르는 선 하나를 긋는다.

    칸은 병합하지 않고 그대로 둔다(결재본과 같은 모양, 담당자 지시). 선은 블록 첫 칸의
    문단에 앵커한 VML 선(o:allowincell)이라 표가 어느 쪽으로 가든 칸을 따라간다 —
    결재본의 선은 쪽 위치에 고정돼 있어 표가 움직이면 어긋났다.
    """
    from docx.table import _Cell
    trs = table._tbl.findall(qn("w:tr"))[first:last + 1]
    width = sum(_width(tc) for tc in trs[0].findall(qn("w:tc")))
    height = 0
    for tr in trs:
        pr = tr.find(qn("w:trPr"))
        h = pr.find(qn("w:trHeight")) if pr is not None else None
        try:
            hv = int(h.get(qn("w:val")))
        except (AttributeError, TypeError, ValueError):
            hv = 0
        height += max(hv, 250)
    w_pt, h_pt = width / 20.0, height / 20.0
    margin = 99 / 20.0                                   # 칸 왼쪽 여백(dxa 99)
    _VML_ID[0] += 1
    sid = _VML_ID[0]
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word">'
        '<w:rPr><w:noProof/></w:rPr><w:pict>'
        f'<v:line id="_x0000_s{sid}" o:spid="_x0000_s{sid}" '
        f'style="position:absolute;z-index:{251650000 + sid};'
        'mso-position-horizontal:absolute;mso-position-horizontal-relative:text;'
        'mso-position-vertical:absolute;mso-position-vertical-relative:text;'
        'mso-wrap-style:square" '
        f'from="{-margin:.2f}pt,{h_pt:.2f}pt" to="{w_pt - margin:.2f}pt,0" '
        f'o:allowincell="t" strokecolor="black" strokeweight="{weight}">'
        '<w10:wrap type="none"/><w10:anchorlock/>'
        '</v:line></w:pict></w:r>'
    )
    from lxml import etree
    run = etree.fromstring(xml)
    first_tc = trs[0].findall(qn("w:tc"))[0]
    p = first_tc.find(qn("w:p"))
    p.append(run)
    for tr in trs:                                       # 블록의 글자는 비운다
        for tc in tr.findall(qn("w:tc")):
            for t in tc.iter(qn("w:t")):
                t.text = ""
            clear_diag(_Cell(tc, table))
    return 1


def keep_paras_before_tables(document):
    """표 바로 앞의 글 있는 문단들에 '다음과 함께' 를 건다.

    '검토 및 승인' 처럼 번호가 없는 제목도 표와 갈라지지 않게 한다(담당자 PC 에서 표지 다음
    쪽의 제목만 앞 쪽에 남는 일이 있었다). 표 앞에 이어진 문단 묶음을 통째로 붙인다.
    """
    body = document.element.body
    kids = list(body)
    n = 0
    for i, el in enumerate(kids):
        if el.tag != qn("w:tbl"):
            continue
        j = i - 1
        while j >= 0 and kids[j].tag == qn("w:p"):
            para = kids[j]
            if para.find(".//" + qn("w:br")) is not None:
                break
            pr = para.find(qn("w:pPr"))
            if pr is not None and pr.find(qn("w:pageBreakBefore")) is not None:
                _keep_next_para(para); n += 1
                break
            _keep_next_para(para); n += 1
            if not "".join(t.text or "" for t in para.iter(qn("w:t"))).strip():
                break                     # 빈 문단까지만 (그 위는 다른 덩어리)
            j -= 1
    return n


def _keep_next_para(para):
    pr = para.find(qn("w:pPr"))
    if pr is None:
        pr = para.makeelement(qn("w:pPr"), {})
        para.insert(0, pr)
    get_or_add(pr, "keepNext")
