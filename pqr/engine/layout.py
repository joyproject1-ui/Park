# -*- coding: utf-8 -*-
"""값을 다 채운 결재본에 회사 조판 규칙을 적용한다 (report-format.md 의 규칙을 코드로).

표는 번호가 아니라 제목·머리행으로 찾는다. 규칙은 어느 제품의 결재본에도 같게 적용된다.
"""
import re
from docx.oxml.ns import qn

from . import docedit as E
from .locate import find_tables, find_para, outline, _text


def _tables_under(document, prefixes):
    out = []
    for p in prefixes:
        out += find_tables(document, p)
    return sorted(set(out))


def _header_text(table):
    return re.sub(r"\s+", "", _text(table._tbl.findall(qn("w:tr"))[0]))


class _SkipFront(Exception):
    """EDMS 서식이라 앞부분 조판을 건너뛴다는 표시."""


def is_edms(document):
    """EDMS 결재본 서식(E-HLF-32)으로 쓴 문서인지 — 바닥글에 EHLF-32 가 있다."""
    for section in document.sections:
        for part in (section.footer, section.header):
            try:
                text = "\n".join(p.text for p in part.paragraphs)
                text += "\n".join(E.cell_text(c) for t in part.tables for r in t.rows for c in E.raw_cells(r))
            except Exception:
                continue
            if re.search(r"E-?HLF-?32", text, re.I):
                return True
    return False


def apply(document, log=None, product_title=None, edms=None):
    """document 를 제자리에서 다듬는다. log(문구) 로 진행을 알린다.

    edms: EDMS 서식이면 True. None 이면 바닥글로 알아낸다. EDMS 서식은 표지·결재표·개정 내역이
    없고 목차 줄 간격도 서식이 정한 것이므로 앞부분 조판을 건너뛴다.
    """
    log = log or (lambda *a: None)
    if edms is None:
        edms = is_edms(document)
    if edms:
        log("EDMS 서식(E-HLF-32) — 앞부분(표지·결재표·개정 내역) 조판 없음")
    T = document.tables
    tables_all = list(range(len(T)))
    approval = 0                                  # 결재표 = 첫 표 (구 분·성 명·직위·서 명·서명일자)
    revision = 1                                  # 개정 내역
    toc = 2                                       # 목차
    for i in range(min(4, len(T))):
        h = _header_text(T[i])
        if "서명일자" in h:
            approval = i
        elif "개정번호" in h:
            revision = i
        elif "목차" in h or "Purpose" in _text(T[i]._tbl):
            toc = i
    # ---- 앞부분 ----
    try:
        if edms:
            raise _SkipFront()
        # 표지 다음은 '검토 및 승인' 이 새 쪽에서 시작한다. 결재본은 표지 아래를 빈 줄로 채워
        # 그 자리를 만들지만, 빈 줄은 글꼴·줄 높이가 조금만 달라도 무너진다(담당자 PC 에서
        # 표지와 결재표가 한 쪽에 붙어 나왔다). 쪽 나눔으로 못 박는다.
        approval_head = find_para(document, "검토 및 승인")
        if approval_head is not None:
            E.page_break_before(approval_head)
        log("개정 내역 위 빈 줄 삭제: %d" % E.drop_blank_paras_after(document, "본 문서의 승인일을 제/개정일로 한다.", 5))
        anchor_title = "3. 대상 제품"
        rev_blank = E.blank_para_before(document, "개 정 내 역", anchor_title)
        E.page_break_before(rev_blank)
        toc_p = find_para(document, "목차 (Table of Contents)")
        if toc_p is not None:
            log("목차 앞 빈 줄 삭제: %d" % E.drop_blank_paras_between(document, T[revision], toc_p))
            toc_blank = E.blank_para_before(document, "목차 (Table of Contents)", anchor_title)
            E.page_break_before(toc_blank)
            E.set_line_spacing(toc_blank, 240, "exact", before=0, after=0)
            E.set_line_spacing(toc_p, 348, "auto", before=0, after=120)
        for row in T[toc].rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    E.set_line_spacing(p, 480, "auto", before=0, after=0)     # 목차 줄 간격 2.0
        p4 = find_para(document, "4. 제품품질평가 일정계획")
        if p4 is not None:
            E.page_break_before(p4)
        E.blank_para_before(document, "5. 책임과 권한", anchor_title)
    except _SkipFront:
        pass
    except Exception as error:                      # 앞부분 서식이 다른 결재본이면 건너뛴다
        log("앞부분 조판 건너뜀: %s" % error)

    # ---- 표 안 정렬 ----
    narrative_tables = set(_tables_under(document, ("11.1", "11.2", "12.")))
    narrative_cols = {}
    for ti in narrative_tables:
        head = [re.sub(r"\s+", "", _text(tc)) for tc in T[ti]._tbl.findall(qn("w:tr"))[0].findall(qn("w:tc"))]
        narrative_cols[ti] = {i for i, h in enumerate(head) if h in ("일탈사항", "조치사항", "변경사항")}
    centered = 0
    for ti, table in enumerate(T):
        if ti == toc:
            continue
        for row in table.rows:
            for ci, cell in enumerate(E.raw_cells(row)):
                if ci in narrative_cols.get(ti, ()):
                    continue
                if E.cell_text(cell).lstrip().startswith("특이사항"):
                    continue
                E.set_cell_valign(cell, "center")
                E.set_cell_align(cell, "center")
                centered += 1
    log("가운데 정렬한 칸: %d" % centered)
    for row in ([] if edms else T[approval].rows):  # 결재표 구분 행은 왼쪽 (EDMS 서식엔 결재표 없음)
        for cell in E.raw_cells(row):
            if any(k in E.cell_text(cell) for k in ("(Written by)", "(Reviewed by)", "(Approved by)")):
                E.set_cell_align(cell, "left")
    for ti in _tables_under(document, ("5.",)):    # 5항 업무 열 왼쪽
        for ri, row in enumerate(T[ti].rows):
            cells = E.raw_cells(row)
            if ri and len(cells) >= 3:
                E.set_cell_align(cells[-1], "left")
    fixed_tables = (toc,) if edms else (approval, revision)   # 서식이 정한 표는 건드리지 않는다
    log("표 안 문단 앞뒤 간격 0: %d" % E.zero_cell_spacing(document, skip=fixed_tables))
    # 칸 폭은 그리드대로 — 단위가 뒤섞인 tcW 를 Word 가 자동 맞춤에 섞어 쓰면 결과가 두 줄로 갈린다
    log("칸 폭을 그리드에 맞춤: %d" % E.fix_all_table_widths(document, skip=(toc,)))
    # 9.2 세부표는 열마다 글 길이가 달라, 한 열만 여러 줄로 늘어지면 줄 수가 같아지도록 다시 나눈다
    # 13.1 도 같이 — 빈 양식의 완료 일자 열이 좁아 '2025.03.2 / 5' 로 잘렸다
    balanced = sum(1 for ti in _tables_under(document, ("9.2", "13.1")) if E.balance_columns(T[ti]) is not None)
    log("9.2·13.1 표 열 폭 균등 배분: %d" % balanced)
    # 표는 모두 '창에 자동으로 맞춤' — 본문 폭을 넘거나 모자라지 않게 (담당자 지시 2026-09)
    page = E.text_width(document)
    log("표 폭을 창에 맞춤: %d" % sum(1 for ti, t in enumerate(T) if ti != toc and E.fit_to_window(t, page)))
    log("표 안 글씨 굴림 10: %d" % E.table_font_size(document, 20, skip=(toc,)))
    from .ooxml_order import get_or_add
    small = 0
    for ti in narrative_tables:                     # 일탈·변경 표 데이터 행은 9 pt
        for tr in T[ti]._tbl.findall(qn("w:tr"))[1:-1]:
            for r in tr.iter(qn("w:r")):
                rpr = r.find(qn("w:rPr"))
                if rpr is None:
                    rpr = r.makeelement(qn("w:rPr"), {}); r.insert(0, rpr)
                for tag in ("sz", "szCs"):
                    get_or_add(rpr, tag).set(qn("w:val"), "18")
                small += 1
    log("일탈·변경 표 9 pt: %d" % small)
    for ti in _tables_under(document, ("9.1",)):   # 허용기준 열 왼쪽
        for ri, row in enumerate(T[ti].rows):
            cells = E.raw_cells(row)
            if ri and len(cells) >= 2:
                E.set_cell_align(cells[-2], "left")

    # ---- 사선 ----
    # '해당 없음' 표: 결재본의 그리기 개체 선을 지우고 빈 블록에 칸 고정 선을 긋는다
    empty_blocks, anchors = {}, []
    for ti, table in enumerate(T):
        trs = table._tbl.findall(qn("w:tr"))
        if len(trs) < 2:
            continue
        body = [i for i, tr in enumerate(trs[1:], 1) if not _text(tr).strip()]
        last_is_comment = _text(trs[-1]).lstrip().startswith("특이사항")
        if body and body == list(range(1, 1 + len(body))) and (last_is_comment or body[-1] == len(trs) - 1):
            empty_blocks[ti] = (body[0], body[-1])
            anchors.append(table._tbl)
    for kind, value, el in outline(document):      # 표 앞 제목 문단에 앵커된 선도
        if kind == "h":
            anchors.append(el)
    log("그리기 개체 선 제거: %d" % E.strip_floating_lines(document, anchors))
    # 내용이 없으면 빈 줄을 하나만 남긴다 (담당자 지시 2026-09) — 그 한 줄에 사선을 긋는다.
    gone = 0
    for ti, (a, b) in list(empty_blocks.items()):
        gone += E.collapse_empty_block(T[ti], a, b)
        empty_blocks[ti] = (a, a)
    log("빈 줄 정리(한 줄만 남김): %d" % gone)
    log("빈 블록 가로지르는 선: %d" % sum(E.draw_block_line(T[ti], a, b) for ti, (a, b) in empty_blocks.items()))
    no_diag = set(empty_blocks)
    log("비고 사선: %d" % sum(E.diag_empty_remarks(t) for ti, t in enumerate(T) if ti not in no_diag))
    slash = 0
    for ti in _tables_under(document, ("9.2",)):
        slash += E.slash_empty_summary(T[ti])
        # 한 표에 요약 블록이 둘(확인 + 수치)인 경우 — '최댓값' 라벨이 둘
        firsts = [i for i, tr in enumerate(T[ti]._tbl.findall(qn("w:tr"))) if _text(tr.findall(qn("w:tc"))[0]).startswith("최댓값")]
        if len(firsts) > 1:
            slash += E.slash_empty_summary(T[ti], rows=T[ti].rows[firsts[0]:firsts[0] + 5])
    log("9.2 요약 사선 덩어리: %d" % slash)
    for ti in _tables_under(document, ("7.",)):    # '기준 / Lot No.' 머리 칸
        for row in T[ti].rows:
            for cell in E.raw_cells(row):
                if "Lot No." in E.cell_text(cell) and "기준" in E.cell_text(cell):
                    E.split_header_cell(cell, "기준", "Lot No.")
    for ti in _tables_under(document, ("8.1.2",)):
        E.diag_empty_in_column(T[ti], "기타업체")
    log("빈 칸 사선: %d" % E.diag_all_empty(document, skip=no_diag | set(fixed_tables) | {toc}))

    # ---- 쪽 배치 ----
    PAGE_BODY = 11800
    whole = 0
    for ti, tbl in enumerate(T):
        if ti == toc or len(tbl.rows) > 21 or E.table_chars(tbl) > 1200:
            continue
        if E.est_height(tbl) > PAGE_BODY * 0.8:
            E._keep_next(tbl.rows[0]); continue
        whole += E.keep_table_together(tbl)
    log("표 통째로 유지 행: %d" % whole)
    for ti, tbl in enumerate(T):
        if ti == toc or len(tbl.rows) <= 21:
            continue
        for r in tbl.rows[:7]:
            E._keep_next(r)
        if E.cell_text(E.raw_cells(tbl.rows[-1])[0]).strip().startswith("Cpk"):
            E.keep_tail_together(tbl, 6)
    protect = ("개 정 내 역", "제품품질평가 보고서") + ((product_title,) if product_title else ())
    log("빈 문단 뭉치 정리: %d" % E.collapse_blank_runs(document, keep=1, min_run=2, protect_before=protect))
    log("각주 뒤 빈 줄 정리: %d" % E.drop_blank_after_note(document))
    log("사선 칸의 N/A 삭제: %d" % E.drop_na_in_diag_cells(document))
    log("항 제목 줄맞춤: %d" % E.align_section_titles(document))
    log("‘확인 필요’ 노랑 표시: %d" % E.highlight(document, "확인 필요"))
    log("행 분할 금지: %d" % E.no_split_rows(document))
    keep = 0
    for ti in _tables_under(document, ("8.2.1", "8.2.2", "13.2")):
        keep += E.keep_merged_groups(T[ti], 3)
    for ti in _tables_under(document, ("9.1", "13.3", "5.")):
        keep += E.keep_merged_groups(T[ti], 0)
    for ti in _tables_under(document, ("13.2",)):
        E.split_year_by_lot(T[ti], 1, 3, 12)
        keep += E.keep_groups(T[ti], 1, 12)
    log("묶음 유지 행: %d" % keep)
    log("페이지 나눔 정리: %d" % E.hard_breaks_to_page_break_before(document))
    p7 = find_para(document, "7. 수율 현황")
    if p7 is not None:
        E.page_break_before(p7)
    h = find_para(document, "13.3.2")
    if h is not None:
        prev = _tables_under(document, ("13.3.1",))
        if prev:
            E.drop_blank_paras_between(document, T[prev[0]], h)
        E.page_break_before(h)
    log("항 제목 다음과 함께: %d" % E.keep_headings_with_next(document))
    log("표 앞 문단 다음과 함께: %d" % E.keep_paras_before_tables(document))
    log("윗첨자 각주 번호: %d" % len(E.superscript_note_marks(document)))
    log("빈 쪽 방지 정리: %s" % E.tidy_page_breaks(document))
    log("항 사이 한 줄 띄움: %d" % E.space_before_sections(document))
    log("각주 내어쓰기: %d" % E.hanging_indent_notes(document))
    try:
        from .toc import link_toc
        log("목차 쪽수 필드: %d" % link_toc(document, T[toc]))
    except Exception as error:
        log("목차 필드 건너뜀: %s" % error)
    return document
