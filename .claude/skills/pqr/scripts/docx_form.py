# -*- coding: utf-8 -*-
"""python-docx 로 회사 서식(병합이 섞인 표)을 안전하게 채우는 헬퍼.

python-docx 의 ``table.cell(r, c)`` 는 모든 행이 그리드 열을 가득 채운다고 가정하고
평면 인덱싱을 합니다. 절차서 서식에는 그리드 열보다 칸이 적은 행이 있어서, 그런 표에서는
행이 내려갈수록 좌표가 밀립니다. 또 세로 병합(vMerge)된 아래쪽 칸에 값을 쓰면 위 칸의
내용이 지워집니다.

여기 있는 함수들은 행마다 ``w:tc`` 를 직접 걸으며 ``gridSpan`` 을 더해 좌표를 찾고,
병합 이어짐 칸에는 쓰지 않습니다.
"""
import copy
from decimal import Decimal, ROUND_HALF_UP

from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def _span(tc):
    pr = tc.find(qn('w:tcPr'))
    if pr is None:
        return 1
    gs = pr.find(qn('w:gridSpan'))
    return int(gs.get(qn('w:val'))) if gs is not None else 1


def _continues_above(tc):
    """세로 병합으로 위 칸에 이어지는 칸이면 True (여기에 쓰면 위 칸이 지워집니다)."""
    pr = tc.find(qn('w:tcPr'))
    if pr is None:
        return False
    vm = pr.find(qn('w:vMerge'))
    return vm is not None and (vm.get(qn('w:val')) or 'continue') == 'continue'


def cell_at(table, row, col):
    """논리적인 (행, 열) 좌표의 ``w:tc``. 해당 칸이 없으면 None."""
    position = 0
    for tc in table.rows[row]._tr.findall(qn('w:tc')):
        width = _span(tc)
        if position <= col < position + width:
            return tc
        position += width
    return None


def _donor_rpr(table):
    """표 안에서 찾은 첫 run 의 글꼴 속성(굵기 제외).

    빈 칸에 add_run 으로 글을 넣으면 run 에 rPr 가 없어 문서 기본 글꼴로
    떨어집니다(이 서식은 본문 굴림 10pt 인데 기본은 바탕). 채운 칸만 글꼴이
    달라 보이는 원인이 이것이라, 같은 표의 기존 run 에서 서식을 빌려 옵니다.
    머리글 run 이 걸리는 일이 많아 굵기(b · bCs)는 떼고 씁니다.
    """
    for tr in table._tbl.iter(qn('w:tr')):
        for r in tr.iter(qn('w:r')):
            rpr = r.find(qn('w:rPr'))
            if rpr is not None:
                donor = copy.deepcopy(rpr)
                for tag in ('w:b', 'w:bCs'):
                    node = donor.find(qn(tag))
                    if node is not None:
                        donor.remove(node)
                return donor
    return None


def set_cell(table, row, col, text):
    """칸의 글을 바꿉니다. 서식(글꼴·정렬)은 첫 run 을 재사용해 유지합니다."""
    tc = cell_at(table, row, col)
    if tc is None or _continues_above(tc):
        return
    cell = _Cell(tc, table)
    para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    if not para.runs:
        run = para.add_run('')
        donor = _donor_rpr(table)
        if donor is not None and run._element.find(qn('w:rPr')) is None:
            run._element.insert(0, donor)
    para.runs[0].text = str(text)
    for run in para.runs[1:]:
        run.text = ''


def clone_row(table, index):
    """``index`` 행을 복제해 바로 아래에 넣고 새 행을 돌려줍니다."""
    row = copy.deepcopy(table.rows[index]._tr)
    table.rows[index]._tr.addnext(row)
    return table.rows[index + 1]


def clone_block(table, start, size, times):
    """``start`` 부터 ``size`` 행을 한 묶음으로 표 끝에 ``times`` 번 덧붙입니다.

    적격성평가 표(설비 한 대 = 문서번호 행 + 완료일 행)처럼 여러 행이 한 항목인 표에 씁니다.
    """
    for _ in range(times):
        anchor = table.rows[len(table.rows) - 1]._tr
        for offset in range(size):
            new = copy.deepcopy(table.rows[start + offset]._tr)
            anchor.addnext(new)
            anchor = new


def fill(table, first_row, records):
    """``records`` 만큼 행을 늘려 채웁니다. 한 record 는 ``[(열, 값), ...]``."""
    for _ in range(len(records) - 1):
        clone_row(table, first_row)
    for offset, record in enumerate(records):
        for col, value in record:
            set_cell(table, first_row + offset, col, value)


def drop_last_row(table):
    tr = table.rows[len(table.rows) - 1]._tr
    tr.getparent().remove(tr)


def set_comment(table, *lines):
    """맨 아래 '특이사항 (Comment)' 행에 줄을 덧붙입니다 (라벨은 그대로 둡니다)."""
    cell = _Cell(cell_at(table, len(table.rows) - 1, 0), table)
    para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    if not para.runs:
        para.add_run('')
    para.runs[0].text = '특이사항 (Comment)'
    for run in para.runs[1:]:
        run.text = ''
    anchor = para._p
    for line in lines:
        new = copy.deepcopy(para._p)
        anchor.addnext(new)
        anchor = new
        target = Paragraph(new, cell)
        target.runs[0].text = line
        for run in target.runs[1:]:
            run.text = ''


def clone_table(table, heading=True):
    """표(와 바로 앞 제목 문단)를 복제해 이어 붙이고 새 Table 을 돌려줍니다.

    제목 문단까지 같이 복제해야 표 둘이 맞붙어 하나로 보이지 않습니다.
    복제한 원본은 아직 채우지 않은 것을 쓰세요 — 이미 행을 지우거나 늘린 표를 복제하면
    구조가 달라집니다.
    """
    new_tbl = copy.deepcopy(table._tbl)
    if heading:
        previous = table._tbl.getprevious()
        if previous is not None:
            table._tbl.addnext(new_tbl)
            table._tbl.addnext(copy.deepcopy(previous))
            return Table(new_tbl, table._parent)
    table._tbl.addnext(new_tbl)
    return Table(new_tbl, table._parent)


def paragraph_after(document, heading, text):
    """``heading`` 으로 시작하는 문단 뒤에 같은 서식의 문단을 넣습니다."""
    for para in document.paragraphs:
        if para.text.strip().startswith(heading):
            new = copy.deepcopy(para._p)
            para._p.addnext(new)
            follower = Paragraph(new, para._parent)
            if not follower.runs:
                follower.add_run('')
            follower.runs[0].text = text
            for run in follower.runs[1:]:
                run.text = ''
            return follower
    return None


def renumber_bookmarks(document):
    """복제한 문단·표가 물고 온 중복 책갈피 ID/이름을 다시 매깁니다.

    문단을 deepcopy 하면 그 안의 책갈피까지 따라오고, 워드 문서는 책갈피 ID 가 유일해야
    스키마 검증을 통과합니다. 저장 직전에 한 번 부르세요.
    """
    counter, pending, names = 0, {}, {}
    for node in document.element.body.iter():
        tag = node.tag.split('}')[1]
        if tag == 'bookmarkStart':
            counter += 1
            pending.setdefault(node.get(qn('w:id')), []).append(str(counter))
            node.set(qn('w:id'), str(counter))
            name = node.get(qn('w:name')) or ''
            names[name] = names.get(name, 0) + 1
            if names[name] > 1:
                node.set(qn('w:name'), '%s_%d' % (name, names[name]))
        elif tag == 'bookmarkEnd':
            queue = pending.get(node.get(qn('w:id')))
            if queue:
                node.set(qn('w:id'), queue.pop(0))


def dump(table):
    """표를 셀 단위로 찍어 봅니다 — 채운 뒤 대조용."""
    for index in range(len(table.rows)):
        cells = []
        for tc in table.rows[index]._tr.findall(qn('w:tc')):
            text = ''.join(n.text or '' for n in tc.iter() if n.tag.endswith('}t'))
            cells.append(' '.join(text.split())[:40])
        print(index, '|', ' | '.join(cells))


# w:tcPr 와 w:tcBorders 의 자식 순서는 스키마로 강제됩니다 — 아무 데나 붙이면
# "This element is not expected" 로 검증에 걸립니다.
_TCPR_ORDER = ('cnfStyle', 'tcW', 'gridSpan', 'hMerge', 'vMerge', 'tcBorders', 'shd',
               'noWrap', 'tcMar', 'textDirection', 'tcFitText', 'vAlign', 'hideMark')
_BORDER_ORDER = ('top', 'start', 'left', 'bottom', 'end', 'right',
                 'insideH', 'insideV', 'tl2br', 'tr2bl')


def _insert_ordered(parent, child, order):
    name = child.tag.split('}')[1]
    rank = order.index(name)
    for existing in parent:
        tag = existing.tag.split('}')[1]
        if tag in order and order.index(tag) > rank:
            existing.addprevious(child)
            return
    parent.append(child)


def close_out(table, row, col):
    """빈 칸에 대각선을 그어 마감 처리합니다 (엑셀 서식과 같은 사내 관행).

    '아직 안 채운 칸'과 '해당 없어 안 채우는 칸'을 구분하는 표시입니다.
    앞으로 채울 칸에는 긋지 마세요 — 시험을 안 한 것처럼 읽힙니다.
    """
    tc = cell_at(table, row, col)
    if tc is None or _continues_above(tc):
        return
    pr = tc.find(qn('w:tcPr'))
    if pr is None:
        pr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, pr)
    borders = pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = pr.makeelement(qn('w:tcBorders'), {})
        _insert_ordered(pr, borders, _TCPR_ORDER)
    existing = borders.find(qn('w:tl2br'))
    if existing is not None:
        borders.remove(existing)
    line = borders.makeelement(qn('w:tl2br'), {})
    line.set(qn('w:val'), 'single')
    line.set(qn('w:sz'), '4')
    line.set(qn('w:color'), '000000')
    _insert_ordered(borders, line, _BORDER_ORDER)


def colliding_columns(table, row, columns):
    """``columns`` 중 같은 칸을 가리키는 조합을 돌려줍니다.

    병합 때문에 논리적으로 다른 두 열이 한 칸으로 해석되면, 나중에 쓴 값이 앞 값을
    소리 없이 덮어씁니다 (빈 값을 쓰면 앞 값이 아예 사라집니다). 서식의 머리글 병합과
    데이터 행의 병합 경계가 어긋나 있을 때 실제로 일어나므로, 표를 채우기 전에 한 번
    확인하세요.

    lxml 프록시는 참조를 놓으면 회수되고 ``id()`` 가 재사용되므로, 여기서는 요소를
    붙들어 둔 채 ``is`` 로 비교합니다.
    """
    held, pairs = {}, []
    for col in columns:
        tc = cell_at(table, row, col)
        if tc is not None:
            held[col] = tc
    keys = list(held)
    for first in range(len(keys)):
        for second in range(first + 1, len(keys)):
            if held[keys[first]] is held[keys[second]]:
                pairs.append((keys[first], keys[second]))
    return pairs


def round_half_up(value, places):
    """사사오입으로 자릿수를 맞춘 문자열.

    파이썬 기본 서식('%.3f')은 짝수 반올림이라 1.0165 를 '1.016' 으로 냅니다.
    보고서 최종본은 '1.017' 이므로 사사오입이 사내 관행입니다.
    """
    quantum = Decimal(1).scaleb(-places)
    return str(Decimal(repr(float(value))).quantize(quantum, rounding=ROUND_HALF_UP))


def summarize(values, places):
    """최댓값·최솟값·평균을 자료와 같은 자릿수로 돌려줍니다.

    평균에 자릿수를 하나 더 쓰고 싶어지지만, 최종본은 자료와 같은 자릿수를 씁니다
    (삼투압 303.67 → '304'). 표 안에서 자릿수가 들쭉날쭉해 보이지 않게 하려는 것입니다.
    """
    numbers = [v for v in values if isinstance(v, (int, float))]
    if not numbers:
        return None
    return tuple(round_half_up(v, places)
                 for v in (max(numbers), min(numbers), sum(numbers) / len(numbers)))
